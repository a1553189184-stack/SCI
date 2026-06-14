from __future__ import annotations

import json
import math
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"


def fmt(x, digits: int = 3) -> str:
    try:
        x = float(x)
        if math.isnan(x):
            return "NA"
        return f"{x:.{digits}f}"
    except Exception:
        return str(x)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
            if i == 0:
                shade(cell, "D9EAF7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_df(doc: Document, title: str, df: pd.DataFrame, cols: list[str], max_rows: int | None = None) -> None:
    doc.add_heading(title, level=2)
    use = df[cols].copy()
    if max_rows:
        use = use.head(max_rows)
    table = doc.add_table(rows=1, cols=len(cols))
    for j, c in enumerate(cols):
        table.rows[0].cells[j].text = c
    for _, row in use.iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(cols):
            val = row[c]
            cells[j].text = fmt(val) if isinstance(val, float) else str(val)
    style_table(table)


def docx_safe_picture(path: Path) -> Path:
    from PIL import Image

    out_dir = ROOT / "outputs_large" / "docx_embedded_images"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{path.stem}.jpg"
    if out.exists() and out.stat().st_mtime >= path.stat().st_mtime:
        return out
    image = Image.open(path).convert("RGB")
    image.save(out, quality=95, optimize=True)
    return out


def add_figure(doc: Document, title: str, path: Path, caption: str, width: float = 6.5) -> None:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(docx_safe_picture(path)), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.style = "Caption"


def make_doc(path: Path, report_mode: bool) -> Path:
    reg = json.loads((ROOT / "outputs_large" / "manuscript_value_registry_large.json").read_text(encoding="utf-8"))
    tables = ROOT / "tables_large"
    figs = ROOT / "figures_large"
    internal = pd.read_csv(tables / "internal_metrics.csv")
    external = pd.read_csv(tables / "external_metrics.csv")
    comp = pd.read_csv(tables / "internal_external_comparison.csv")
    cal = pd.read_csv(tables / "calibration_metrics.csv")
    drop = pd.read_csv(tables / "performance_drop.csv")
    split = pd.read_csv(ROOT / "splits_large" / "hf_large_nih_split_statistics.csv")
    t3 = pd.read_csv(tables / "table3_internal_external_performance_with_prevalence_ci.csv")
    macro_comp_path = tables / "table3_model_macro_comparison.csv"
    macro_comp = pd.read_csv(macro_comp_path) if macro_comp_path.exists() else pd.DataFrame()
    subgroup = pd.read_csv(tables / "table5_subgroup_analysis.csv")
    label_map = pd.read_csv(tables / "table2_label_harmonization.csv")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    title = (
        "Route B Large Public-Dataset Validation Report: NIH ChestX-ray14 to VinDr-CXR/VinBigData-derived External Validation"
        if report_mode
        else "Cross-dataset Generalizability, Calibration, and Explainability of DenseNet121 with ResNet50 and EfficientNet-B0 Comparison for Multi-label Chest Radiograph Diagnosis: Expanded NIH ChestX-ray14 to VinDr-CXR/VinBigData-derived External Validation"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph(
        "Evidence source: real public Hugging Face NIH ChestX-ray14 parquet subset and real public VinDr-CXR/VinBigData-derived PNG subset. "
        "No synthetic metrics were used. Thresholds were derived only from NIH validation. Calibration models were fitted only on NIH calibration."
    )

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"The upgraded Route B experiment expanded the development data to {reg['dataset']['nih_images']} NIH images from {reg['dataset']['nih_patients']} patients and expanded external validation to {reg['dataset']['external_images']} VinDr-derived images. "
        f"DenseNet121 completed {reg['training']['epochs_completed']} epochs; the best checkpoint was selected at epoch {reg['training']['best_epoch']} by validation macro AUPRC ({reg['training']['best_val_macro_auprc']}). "
        f"Internal macro AUROC/AUPRC were {reg['performance']['internal_macro_auroc']}/{reg['performance']['internal_macro_auprc']}; external macro AUROC/AUPRC were {reg['performance']['external_macro_auroc']}/{reg['performance']['external_macro_auprc']}."
    )
    if not macro_comp.empty:
        comparator_text = []
        for _, row in macro_comp[macro_comp["model"].astype(str) != "DenseNet121"].iterrows():
            comparator_text.append(
                f"{row['model']} internal macro AUROC/AUPRC were {fmt(row['internal_macro_auroc'])}/{fmt(row['internal_macro_auprc'])}, "
                f"and external macro AUROC/AUPRC were {fmt(row['external_macro_auroc'])}/{fmt(row['external_macro_auprc'])}"
            )
        if comparator_text:
            doc.add_paragraph(
                "Architecture comparators were trained under the same NIH split and validation-threshold rule. "
                + "; ".join(comparator_text)
                + ". These comparator results are reported as model-comparison evidence rather than replacing the prespecified DenseNet121 main analysis."
            )

    doc.add_heading("Data Expansion Plan and Execution", level=1)
    for text in [
        "NIH expansion was performed by downloading multiple public Hugging Face parquet shards, extracting image bytes to PNG files, harmonizing target labels, and creating patient-level train/validation/calibration/internal-test splits.",
        "VinDr expansion was performed from a public VinDr-CXR/VinBigData-derived PNG mirror. Images were selected from available cached public files with enrichment for externally mappable labels and No Finding. VinDr images remained fully independent from NIH training, validation, threshold selection, and calibration fitting.",
        "The large-run project structure uses metadata_large/, splits_large/, predictions_large/, tables_large/, figures_large/, models_large/, logs_large/, and outputs_large/ so that the original small public-subset experiment remains intact as supplementary evidence.",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("Methods to Update Manuscript", level=1)
    methods = [
        "Title: describe the study as an expanded public-subset NIH to VinDr-CXR/VinBigData-derived external validation, not as a full official-cohort validation.",
        "Abstract: replace small-subset values with large-run sample sizes, macro AUROC/AUPRC, calibration metrics, and Grad-CAM++ case count.",
        "Methods: add Hugging Face parquet/PNG extraction details, patient-level splitting, validation-only thresholds, NIH-only calibration fitting, and external independence rules.",
        "Results: use Table 1-5 and Figure 1-5 from tables_large/ and figures_large/ only.",
        "Discussion: emphasize dataset shift, label-source heterogeneity, class imbalance, and calibration transfer limits.",
        "Limitations: explicitly state that Edema and Pneumonia are unavailable as external labels in the VinDr-derived subset and should not support external primary claims.",
        "Conclusion: keep claims limited to public-dataset validation and reproducible pipeline evidence; do not claim clinical readiness.",
    ]
    for item in methods:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        f"Internal macro AUROC was {reg['performance']['internal_macro_auroc']} and macro AUPRC was {reg['performance']['internal_macro_auprc']}. "
        f"External macro AUROC was {reg['performance']['external_macro_auroc']} and macro AUPRC was {reg['performance']['external_macro_auprc']}. "
        f"The macro AUROC internal-minus-external difference was {reg['performance']['macro_auroc_absolute_drop']}, indicating no observed AUROC degradation in this selected public external subset."
    )
    doc.add_paragraph(
        f"Internal uncalibrated macro Brier score was {reg['calibration']['internal_uncalibrated_macro_brier']}, ECE was {reg['calibration']['internal_uncalibrated_macro_ece']}, and MCE was {reg['calibration']['internal_uncalibrated_macro_mce']}. "
        f"Temperature scaling produced internal macro ECE {reg['calibration']['internal_temperature_macro_ece']} and external macro ECE {reg['calibration']['external_temperature_macro_ece']}. "
        "ECE improvement should be interpreted together with slope and intercept because Platt or isotonic methods can reduce bin-wise error while creating unstable probability-scale slopes in sparse or imbalanced labels."
    )
    doc.add_paragraph(
        f"Grad-CAM++ generated {reg['gradcam']['n_cases']} TP/FP/FN/TN examples across NIH internal and external validation for failure-mode review. These maps are not used as proof of clinical reasoning."
    )

    add_df(doc, "Table 1. Dataset Characteristics", pd.concat([split[["split", "n_patients", "n_images"]], pd.DataFrame([{"split": "VinDr-derived external", "n_patients": reg["dataset"]["external_patients"], "n_images": reg["dataset"]["external_images"]}])]), ["split", "n_patients", "n_images"])
    add_df(doc, "Table 2. Label Harmonization", label_map, ["harmonized_label", "nih_terms", "vindr_terms", "role"], max_rows=8)
    summary_t3 = t3[t3["average"] == "label"].copy()
    summary_t3_display = pd.DataFrame(
        {
            "model": summary_t3["model"],
            "dataset": summary_t3["dataset_group"].replace({"NIH internal test": "NIH internal", "VinDr-derived external": "VinDr external"}),
            "label": summary_t3["label"],
            "prevalence": summary_t3["prevalence"].map(lambda x: fmt(x)),
            "AUROC (95% CI)": summary_t3.apply(lambda r: f"{fmt(r['auroc'])} ({fmt(r['auroc_ci_lower'])}-{fmt(r['auroc_ci_upper'])})", axis=1),
            "AUPRC (95% CI)": summary_t3.apply(lambda r: f"{fmt(r['auprc'])} ({fmt(r['auprc_ci_lower'])}-{fmt(r['auprc_ci_upper'])})", axis=1),
            "F1": summary_t3.apply(lambda r: "NA" if pd.isna(r["auroc"]) or float(r["prevalence"]) == 0.0 else fmt(r["f1"]), axis=1),
        }
    )
    add_df(doc, "Table 3. Internal and External Performance by Label and Model with Prevalence and 95% CI", summary_t3_display, list(summary_t3_display.columns), max_rows=48)
    if not macro_comp.empty:
        macro_display = pd.DataFrame(
            {
                "model": macro_comp["model"],
                "internal AUROC": macro_comp["internal_macro_auroc"].map(lambda x: fmt(x)),
                "internal AUPRC": macro_comp["internal_macro_auprc"].map(lambda x: fmt(x)),
                "external AUROC": macro_comp["external_macro_auroc"].map(lambda x: fmt(x)),
                "external AUPRC": macro_comp["external_macro_auprc"].map(lambda x: fmt(x)),
                "AUROC drop": macro_comp["macro_auroc_absolute_drop"].map(lambda x: fmt(x)),
                "AUPRC drop": macro_comp["macro_auprc_absolute_drop"].map(lambda x: fmt(x)),
            }
        )
        add_df(doc, "Table 3a. Macro-level Model Comparison", macro_display, list(macro_display.columns))
    table4 = cal[(cal["average"] == "macro")][["dataset", "method", "brier", "ece", "mce", "slope", "intercept"]]
    add_df(doc, "Table 4. Calibration Metrics", table4, list(table4.columns))
    add_df(doc, "Table 5. Sensitivity and Subgroup Analysis", subgroup, ["subgroup_variable", "subgroup", "n_images", "n_patients", "macro_auroc", "macro_auprc", "macro_f1", "macro_accuracy"])

    doc.add_page_break()
    add_figure(doc, "Figure 1. Study Workflow", figs / "figure1_study_workflow_large.png", "Expanded NIH development workflow with independent VinDr-derived external validation.")
    add_figure(doc, "Figure 2. Internal vs External AUROC/AUPRC", figs / "figure2_internal_vs_external_auroc_auprc_ci.png", "Label-wise internal and external discrimination metrics.")
    add_figure(doc, "Figure 3. Calibration Curves", figs / "figure3_calibration_curves_large.png", "Reliability diagrams after NIH-only calibration fitting and transfer.")
    add_figure(doc, "Figure 4. Grad-CAM++ Failure-mode Examples", figs / "figure4_gradcam_failure_modes_large.png", "Representative TP/FP/FN/TN heatmap panels.")
    add_figure(doc, "Figure 5. External Performance Drop by Label", figs / "figure5_external_performance_drop_by_label.png", "Internal minus external AUROC; negative values indicate higher AUROC in the external subset.")
    if (figs / "supplementary_figure_original_small_subset.png").exists():
        add_figure(doc, "Supplementary Figure. Original Small Public-subset Result", figs / "supplementary_figure_original_small_subset.png", "Small public-subset result retained as supplementary pipeline history.")

    doc.add_heading("Reviewer-risk Audit", level=1)
    risks = [
        "Subset rather than full official release: still present; mitigated by explicitly describing public executable subsets.",
        "Limited sample size: improved from 1229/494 to 5000/1000 images, but still not full cohort.",
        "Incomplete VinDr labels for Edema or Pneumonia: still present; these labels are excluded from external primary claims.",
        "Weak NIH report-mined labels: still present; discussed as label-noise limitation.",
        "Class imbalance: still present; handled with positive class weights and prevalence reporting.",
        "Confidence intervals: addressed with patient-level bootstrap CI.",
        "Calibration overclaiming: mitigated by reporting Brier, ECE, MCE, slope, intercept, and caution about unstable post-hoc calibration.",
        "Grad-CAM overinterpretation: mitigated by explicitly limiting Grad-CAM++ to failure-mode analysis.",
    ]
    for item in risks:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Re-run Commands", level=1)
    commands = [
        r".\.venv\Scripts\python.exe scripts\prepare_hf_large.py --nih-max-images 5000 --nih-train-shards 5 --vindr-target-images 1000 --epochs 5",
        r".\.venv\Scripts\python.exe scripts\train.py --config configs\hf_large.yaml",
        r".\.venv\Scripts\python.exe scripts\evaluate_internal.py --config configs\hf_large.yaml",
        r".\.venv\Scripts\python.exe scripts\evaluate_external.py --config configs\hf_large.yaml",
        r".\.venv\Scripts\python.exe scripts\calibration.py --config configs\hf_large.yaml",
        r".\.venv\Scripts\python.exe scripts\gradcam.py --config configs\hf_large.yaml --max-cases-per-label 1",
        r".\.venv\Scripts\python.exe scripts\make_resnet50_config.py",
        r".\.venv\Scripts\python.exe scripts\train.py --config configs\hf_large_resnet50.yaml",
        r".\.venv\Scripts\python.exe scripts\evaluate_internal.py --config configs\hf_large_resnet50.yaml",
        r".\.venv\Scripts\python.exe scripts\evaluate_external.py --config configs\hf_large_resnet50.yaml",
        r".\.venv\Scripts\python.exe scripts\calibration.py --config configs\hf_large_resnet50.yaml",
        r".\.venv\Scripts\python.exe scripts\gradcam.py --config configs\hf_large_resnet50.yaml --max-cases-per-label 1",
        r".\.venv\Scripts\python.exe scripts\make_efficientnet_b0_config.py",
        r".\.venv\Scripts\python.exe scripts\train.py --config configs\hf_large_efficientnet_b0.yaml",
        r".\.venv\Scripts\python.exe scripts\evaluate_internal.py --config configs\hf_large_efficientnet_b0.yaml",
        r".\.venv\Scripts\python.exe scripts\evaluate_external.py --config configs\hf_large_efficientnet_b0.yaml",
        r".\.venv\Scripts\python.exe scripts\calibration.py --config configs\hf_large_efficientnet_b0.yaml",
        r".\.venv\Scripts\python.exe scripts\gradcam.py --config configs\hf_large_efficientnet_b0.yaml --max-cases-per-label 1",
        r".\.venv\Scripts\python.exe scripts\build_large_assets.py",
    ]
    for cmd in commands:
        doc.add_paragraph(cmd, style="Intense Quote")

    doc.save(path)
    return path


if __name__ == "__main__":
    report = make_doc(DESKTOP / "CXR_RouteB_Large_Experiment_Report.docx", report_mode=True)
    manuscript = make_doc(DESKTOP / "CXR_Manuscript_Version_C_RouteB_Large_COMPLETED.docx", report_mode=False)
    shutil.copy2(report, ROOT / "outputs_large" / report.name)
    shutil.copy2(manuscript, ROOT / "outputs_large" / manuscript.name)
    print(report)
    print(manuscript)


