from __future__ import annotations

import json
import math
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"


def fmt(x, digits: int = 3) -> str:
    try:
        x = float(x)
        if math.isnan(x):
            return "NA"
        return f"{x:.{digits}f}"
    except Exception:
        return str(x)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_df_table(doc: Document, df: pd.DataFrame, columns: list[str], title: str, note: str | None = None, max_rows: int | None = None) -> None:
    doc.add_paragraph(title, style="Heading 2")
    table_df = df[columns].copy()
    if max_rows is not None:
        table_df = table_df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(columns))
    hdr = table.rows[0].cells
    for j, col in enumerate(columns):
        hdr[j].text = col
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                value = fmt(value)
            cells[j].text = str(value)
    style_table(table)
    if note:
        p = doc.add_paragraph(note)
        p.style = "Intense Quote"


def add_figure(doc: Document, path: Path, title: str, caption: str, width: float = 6.4) -> None:
    doc.add_paragraph(title, style="Heading 2")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.style = "Caption"


def metric_rows() -> dict:
    reg = json.loads((ROOT / "outputs" / "manuscript_value_registry.json").read_text(encoding="utf-8"))
    internal = pd.read_csv(ROOT / "tables" / "internal_metrics.csv")
    external = pd.read_csv(ROOT / "tables" / "external_metrics.csv")
    comparison = pd.read_csv(ROOT / "tables" / "internal_external_comparison.csv")
    cal = pd.read_csv(ROOT / "tables" / "calibration_metrics.csv")
    return {"reg": reg, "internal": internal, "external": external, "comparison": comparison, "cal": cal}


def build_docx() -> Path:
    data = metric_rows()
    reg = data["reg"]
    perf = reg["performance"]
    calreg = reg["calibration"]
    ds = reg["dataset"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"

    title = "Cross-dataset Generalizability, Calibration, and Explainability of DenseNet121 for Multi-label Chest Radiograph Diagnosis: NIH ChestX-ray14 to VinDr-CXR Public Subset Validation"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Completed Version C - evidence-aligned manuscript draft").italic = True

    doc.add_paragraph(
        "This version is based only on real prediction CSV files and generated figures from the local reproducible pipeline. "
        "The analysis uses a publicly accessible Hugging Face NIH ChestX-ray14 parquet subset for development and a VinDr-CXR/VinBigData-derived public PNG subset for external validation. "
        "Because this executable subset is smaller than the full public datasets, all performance claims are framed as public-subset evidence rather than full-cohort clinical validation."
    )

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Background: Chest radiograph deep learning models are commonly evaluated using internal discrimination metrics, but cross-dataset generalizability, probability calibration, and error-mode visualization are essential before broad clinical claims can be made."
    )
    doc.add_paragraph(
        "Purpose: To evaluate a reproducible DenseNet121 pipeline for multi-label chest radiograph classification using NIH ChestX-ray14 as the development dataset and a VinDr-CXR/VinBigData-derived public subset as independent external validation."
    )
    doc.add_paragraph(
        f"Materials and Methods: NIH data included {ds['nih_images']} images from {ds['nih_patients']} patients and were split at patient level into training, validation, calibration, and internal test subsets. "
        f"The external validation set included {ds['vindr_images']} images. DenseNet121 was trained with BCEWithLogitsLoss, positive class weights, AdamW, mixed precision, and validation macro AUPRC model selection. "
        "Thresholds were selected only on the NIH validation split. Calibration models were fitted only on the NIH calibration split. Grad-CAM++ was used for failure-mode analysis."
    )
    doc.add_paragraph(
        f"Results: Internal testing yielded macro AUROC {perf['internal_macro_auroc']} and macro AUPRC {perf['internal_macro_auprc']}. "
        f"External validation yielded macro AUROC {perf['external_macro_auroc']} and macro AUPRC {perf['external_macro_auprc']}. "
        f"The macro AUROC absolute drop was {perf['macro_auroc_absolute_drop']} (relative drop {perf['macro_auroc_relative_drop']}). "
        f"Uncalibrated internal test calibration showed macro Brier score {calreg['internal_uncalibrated_macro_brier']}, ECE {calreg['internal_uncalibrated_macro_ece']}, MCE {calreg['internal_uncalibrated_macro_mce']}, slope {calreg['internal_uncalibrated_macro_slope']}, and intercept {calreg['internal_uncalibrated_macro_intercept']}. "
        f"Temperature scaling reduced internal macro ECE to {calreg['internal_temperature_macro_ece']} and external macro ECE to {calreg['external_temperature_macro_ece']}. "
        f"Grad-CAM++ generated {reg['gradcam']['n_cases']} TP/FP/FN/TN case panels for error-mode review."
    )
    doc.add_paragraph(
        "Conclusion: In this public-subset experiment, discrimination transferred only modestly across datasets, calibration changed after post-hoc scaling, and Grad-CAM++ highlighted heterogeneous failure modes. The results support careful external validation and calibration reporting, while avoiding claims of clinical readiness."
    )
    doc.add_paragraph("Keywords: chest radiography; deep learning; external validation; calibration; Grad-CAM; public datasets; multi-label classification")

    doc.add_heading("Introduction", level=1)
    intro = [
        "Public chest radiograph datasets have enabled reproducible development of deep learning models, but internal testing alone can overstate real-world reliability when training and test data share acquisition protocols, label-generation processes, and patient case mix.",
        "Cross-dataset validation is particularly important for chest radiograph AI because findings are heterogeneous, labels may be report-mined or radiologist-annotated, and disease prevalence can shift substantially across institutions and datasets.",
        "Discrimination metrics such as AUROC and AUPRC quantify ranking performance but do not establish that predicted probabilities are numerically meaningful. Calibration analysis is therefore required when outputs may be interpreted as risk estimates or used with decision thresholds.",
        "Explainability maps can provide useful failure-mode review, but heatmaps should not be interpreted as proof that a model uses clinically valid reasoning. In this study, Grad-CAM++ is used only to inspect TP, FP, FN, and TN examples."
    ]
    for paragraph in intro:
        doc.add_paragraph(paragraph)

    doc.add_heading("Materials and Methods", level=1)
    doc.add_heading("Study Design and Data Sources", level=2)
    doc.add_paragraph(
        f"This retrospective computational experiment used NIH ChestX-ray14 as the development source. A real public Hugging Face parquet shard was extracted into image files and metadata, yielding {ds['nih_images']} images from {ds['nih_patients']} patients. "
        f"External validation used {ds['vindr_images']} images from a public VinDr-CXR/VinBigData-derived PNG subset. The external set was never used for training, threshold selection, model selection, or calibration fitting."
    )
    doc.add_heading("Labels", level=2)
    doc.add_paragraph(
        "The primary harmonized labels were Atelectasis, Cardiomegaly, Pleural Effusion, Pneumothorax, and Consolidation. Edema and Pneumonia were retained in the model output because they are present in the NIH label map, but Edema and Pneumonia were not available as source labels in the VinDr-derived external subset and were not used for external primary conclusions. No Finding was treated as a sensitivity label."
    )
    doc.add_heading("Model and Training", level=2)
    doc.add_paragraph(
        f"DenseNet121 was trained at 224 x 224 input resolution for {reg['training']['epochs_completed']} epochs. The best checkpoint was selected by validation macro AUPRC ({reg['training']['best_val_macro_auprc']}). "
        "The loss was BCEWithLogitsLoss with positive class weights; optimization used AdamW and mixed precision when CUDA was available."
    )
    doc.add_heading("Evaluation and Calibration", level=2)
    doc.add_paragraph(
        "Internal performance was evaluated on the NIH internal test split. External performance was evaluated by applying the same checkpoint directly to the VinDr-derived external subset. Binary thresholds were selected only from the NIH validation split. Calibration models were fitted only on the NIH calibration split and then applied to NIH internal test and external predictions."
    )
    doc.add_heading("Explainability", level=2)
    doc.add_paragraph(
        "Grad-CAM++ was generated for selected TP, FP, FN, and TN examples. These heatmaps were used as visual failure-mode analysis and not as evidence of clinically valid reasoning."
    )

    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        f"The NIH development subset contained {ds['nih_images']} images from {ds['nih_patients']} patients. The patient-level split assigned 789 images to training, 131 to validation, 154 to calibration, and 155 to internal testing. The external validation subset contained {ds['vindr_images']} images."
    )
    doc.add_paragraph(
        f"Internal macro AUROC was {perf['internal_macro_auroc']} and macro AUPRC was {perf['internal_macro_auprc']}. External macro AUROC was {perf['external_macro_auroc']} and macro AUPRC was {perf['external_macro_auprc']}. The macro AUROC decrease from internal to external evaluation was {perf['macro_auroc_absolute_drop']}, corresponding to a relative decrease of {perf['macro_auroc_relative_drop']}."
    )
    doc.add_paragraph(
        f"Calibration analysis showed internal uncalibrated macro Brier score {calreg['internal_uncalibrated_macro_brier']} and ECE {calreg['internal_uncalibrated_macro_ece']}. Temperature scaling fitted on the NIH calibration split reduced internal macro ECE to {calreg['internal_temperature_macro_ece']} and external macro ECE to {calreg['external_temperature_macro_ece']}. Calibration slopes remained below 1 in the macro summary, indicating residual probability-scale miscalibration."
    )
    doc.add_paragraph(
        "Grad-CAM++ panels showed heterogeneous activation patterns across TP, FP, FN, and TN examples. The panels are most useful for identifying possible shortcut reliance or missed disease regions, but they do not establish causal model reasoning."
    )

    # Tables.
    split = pd.read_csv(ROOT / "splits" / "hf_subset_real_nih_split_statistics.csv")
    ext = pd.DataFrame([{"split": "VinDr external", "n_patients": ds["vindr_patients"], "n_images": ds["vindr_images"]}])
    table1 = pd.concat([split[["split", "n_patients", "n_images"]], ext], ignore_index=True)
    add_df_table(doc, table1, ["split", "n_patients", "n_images"], "Table 1. Dataset characteristics and split sizes")

    table2 = pd.read_csv(ROOT / "tables" / "table2_label_harmonization.csv")
    add_df_table(doc, table2, ["harmonized_label", "nih_terms", "vindr_terms", "role"], "Table 2. Label harmonization", max_rows=8)

    comp = pd.read_csv(ROOT / "tables" / "internal_external_comparison.csv")
    table3 = comp[comp["average"] == "label"].copy()
    table3 = table3[["label", "internal_auroc", "internal_auprc", "external_auroc", "external_auprc"]]
    for col in table3.columns[1:]:
        table3[col] = table3[col].map(fmt)
    add_df_table(doc, table3, list(table3.columns), "Table 3. Internal and external performance by label")

    calm = pd.read_csv(ROOT / "tables" / "calibration_metrics.csv")
    table4 = calm[(calm["average"] == "macro") & (calm["method"].isin(["uncalibrated", "temperature", "platt", "isotonic"]))].copy()
    table4 = table4[["dataset", "method", "brier", "ece", "mce", "slope", "intercept"]]
    for col in ["brier", "ece", "mce", "slope", "intercept"]:
        table4[col] = table4[col].map(fmt)
    add_df_table(doc, table4, list(table4.columns), "Table 4. Calibration metrics before and after calibration")

    sub = pd.read_csv(ROOT / "tables" / "table5_subgroup_analysis.csv")
    for col in ["macro_auroc", "macro_auprc", "macro_f1", "macro_accuracy"]:
        sub[col] = sub[col].map(fmt)
    add_df_table(doc, sub, ["subgroup_variable", "subgroup", "n_images", "n_patients", "macro_auroc", "macro_auprc", "macro_f1", "macro_accuracy"], "Table 5. Sensitivity and subgroup analysis")

    # Figures.
    doc.add_page_break()
    add_figure(doc, ROOT / "figures" / "figure1_study_workflow.png", "Figure 1. Study workflow", "Patient-level NIH development splits were used for training, validation threshold selection, calibration fitting, and internal testing; external validation used the VinDr-derived subset without retraining or tuning.")
    add_figure(doc, ROOT / "figures" / "figure2_internal_vs_external_auroc_auprc.png", "Figure 2. Internal versus external AUROC/AUPRC", "Label-wise discrimination metrics comparing NIH internal test and VinDr-derived external validation.")
    add_figure(doc, ROOT / "figures" / "figure3_calibration_curves.png", "Figure 3. Calibration curves", "Reliability diagrams for NIH internal test and VinDr-derived external validation after applying calibration models fitted only on the NIH calibration split.")
    add_figure(doc, ROOT / "figures" / "figure4_gradcam_examples.png", "Figure 4. Grad-CAM++ failure-mode examples", "Representative TP, FP, FN, and TN Grad-CAM++ examples from NIH internal and VinDr-derived external predictions.")
    add_figure(doc, ROOT / "figures" / "figure5_external_performance_drop_by_label.png", "Figure 5. External performance drop by label", "Internal minus external AUROC by label. Positive values indicate lower external AUROC.")
    add_figure(doc, ROOT / "figures" / "supplementary_figure1_medmnist_pipeline_verification.png", "Supplementary Figure 1. MedMNIST pipeline verification", "The MedMNIST experiment is retained only as supplementary software verification and is not used for the main CXR claims.")

    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(
        "This executable public-subset experiment demonstrates that a complete CXR generalizability and calibration pipeline can be run end to end with patient-level splitting, validation-derived thresholds, internal calibration fitting, independent external prediction, and heatmap-based error review. The observed internal and external macro AUROC values were similar, but label-wise AUPRC and threshold metrics varied substantially, reinforcing the need to report more than AUROC."
    )
    doc.add_paragraph(
        "Calibration improved for some summaries after temperature scaling, but calibration remained imperfect after transfer to the external subset. This supports reporting calibration metrics and reliability diagrams whenever model outputs may be interpreted as probabilities."
    )
    doc.add_paragraph(
        "The main limitation is that the present completed run used a public executable subset rather than the full NIH ChestX-ray14 and official VinDr-CXR DICOM releases. The external source is VinDr-CXR/VinBigData-derived PNG data, and Edema and Pneumonia were not available as external source labels. The results should therefore be interpreted as evidence for the pipeline and a public-subset validation experiment, not as a definitive full-dataset clinical validation study."
    )
    doc.add_paragraph(
        "Grad-CAM++ results should be treated as qualitative failure-mode review only. Heatmaps can suggest where the model was sensitive, but they cannot prove that the network used clinically appropriate features or exclude shortcut learning."
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "A reproducible DenseNet121 chest radiograph pipeline was completed using real public NIH and VinDr-derived data subsets. The experiment produced internal and external predictions, performance tables, calibration analysis, reliability diagrams, Grad-CAM++ panels, and a reproducibility package. The evidence supports cautious public-dataset claims about generalizability and calibration transfer, while reserving full clinical claims for larger official full-cohort validation."
    )

    doc.add_heading("Data and Code Availability", level=1)
    doc.add_paragraph(
        "All local outputs used for this manuscript are stored in the project reproducibility package. Public source data were accessed through Hugging Face mirrors of NIH ChestX-ray14 and VinDr-CXR/VinBigData-derived PNG data. The code records random seed, package versions, GPU information, split CSV files, predictions, calibration outputs, figures, logs, and checkpoints."
    )

    doc.add_heading("Ethics Statement", level=1)
    doc.add_paragraph(
        "This study used publicly available de-identified imaging datasets and did not involve direct interaction with human participants. Institutional review requirements depend on local policy for secondary analysis of public de-identified data."
    )

    doc.add_heading("References", level=1)
    refs = [
        "Wang X, Peng Y, Lu L, Lu Z, Bagheri M, Summers RM. ChestX-ray8: hospital-scale chest X-ray database and benchmarks on weakly-supervised classification and localization of common thorax diseases. IEEE CVPR. 2017.",
        "Nguyen HQ, Lam K, Le LT, et al. VinDr-CXR: An open dataset of chest X-rays with radiologist's annotations. Scientific Data. 2022;9:429.",
        "Nguyen HQ, Pham HH, Le TL, Dao M, Lam K. VinDr-CXR: An open dataset of chest X-rays with radiologist annotations (version 1.0.0). PhysioNet. 2021. doi:10.13026/3akn-b287.",
        "Huang G, Liu Z, van der Maaten L, Weinberger KQ. Densely connected convolutional networks. IEEE CVPR. 2017.",
        "Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. ICML. 2017.",
        "Selvaraju RR, Cogswell M, Das A, Vedantam R, Parikh D, Batra D. Grad-CAM: visual explanations from deep networks via gradient-based localization. ICCV. 2017.",
        "Chattopadhay A, Sarkar A, Howlader P, Balasubramanian VN. Grad-CAM++: generalized gradient-based visual explanations for deep convolutional networks. WACV. 2018.",
    ]
    for item in refs:
        doc.add_paragraph(item, style="List Number")

    out = DESKTOP / "CXR_Manuscript_Version_C_NIH_to_VinDr_COMPLETED_HF_Subset.docx"
    doc.save(out)
    shutil.copy2(out, ROOT / "outputs" / out.name)
    return out


if __name__ == "__main__":
    print(build_docx())


