from __future__ import annotations

import json
import math
import shutil
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from build_routeb_large_report import add_df, add_figure, fmt  # noqa: E402

DESKTOP = Path.home() / "Desktop"


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_section_paragraphs(doc: Document, heading: str, paragraphs: list[str]) -> None:
    doc.add_heading(heading, level=1)
    for text in paragraphs:
        add_para(doc, text)


def compact_table3(t3: pd.DataFrame) -> pd.DataFrame:
    rows = t3[t3["average"] == "label"].copy()
    return pd.DataFrame(
        {
            "model": rows["model"],
            "dataset": rows["dataset_group"].replace({"NIH internal test": "NIH internal", "VinDr-derived external": "VinDr external"}),
            "label": rows["label"],
            "prevalence": rows["prevalence"].map(lambda x: fmt(x)),
            "AUROC (95% CI)": rows.apply(lambda r: f"{fmt(r['auroc'])} ({fmt(r['auroc_ci_lower'])}-{fmt(r['auroc_ci_upper'])})", axis=1),
            "AUPRC (95% CI)": rows.apply(lambda r: f"{fmt(r['auprc'])} ({fmt(r['auprc_ci_lower'])}-{fmt(r['auprc_ci_upper'])})", axis=1),
            "F1": rows.apply(lambda r: "NA" if pd.isna(r["auroc"]) or float(r["prevalence"]) == 0.0 else fmt(r["f1"]), axis=1),
        }
    )


def macro_display(macro: pd.DataFrame) -> pd.DataFrame:
    return pd.DataFrame(
        {
            "model": macro["model"],
            "internal AUROC": macro["internal_macro_auroc"].map(lambda x: fmt(x)),
            "internal AUPRC": macro["internal_macro_auprc"].map(lambda x: fmt(x)),
            "external AUROC": macro["external_macro_auroc"].map(lambda x: fmt(x)),
            "external AUPRC": macro["external_macro_auprc"].map(lambda x: fmt(x)),
            "AUROC drop": macro["macro_auroc_absolute_drop"].map(lambda x: fmt(x)),
            "AUPRC drop": macro["macro_auprc_absolute_drop"].map(lambda x: fmt(x)),
        }
    )


def make_manuscript(path: Path) -> Path:
    tables = ROOT / "tables_large"
    figs = ROOT / "figures_large"
    reg = json.loads((ROOT / "outputs_large" / "manuscript_value_registry_large.json").read_text(encoding="utf-8"))

    split = pd.read_csv(ROOT / "splits_large" / "hf_large_nih_split_statistics.csv")
    table1 = pd.concat(
        [
            split[["split", "n_patients", "n_images"]],
            pd.DataFrame([{"split": "VinDr-derived external", "n_patients": reg["dataset"]["external_patients"], "n_images": reg["dataset"]["external_images"]}]),
        ],
        ignore_index=True,
    )
    label_map = pd.read_csv(tables / "table2_label_harmonization.csv")
    t3 = pd.read_csv(tables / "table3_internal_external_performance_with_prevalence_ci.csv")
    macro = pd.read_csv(tables / "table3_model_macro_comparison.csv")
    cal = pd.read_csv(tables / "calibration_metrics.csv")
    subgroup = pd.read_csv(tables / "table5_subgroup_analysis.csv")

    comparators = macro[macro["model"].astype(str) != "DenseNet121"].copy()
    comparator_sentence = "; ".join(
        [
            f"{row['model']} internal macro AUROC/AUPRC {fmt(row['internal_macro_auroc'])}/{fmt(row['internal_macro_auprc'])} and external macro AUROC/AUPRC {fmt(row['external_macro_auroc'])}/{fmt(row['external_macro_auprc'])}"
            for _, row in comparators.iterrows()
        ]
    )

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    title = "Generalizability, Calibration, and Explainability of Deep Learning Models for Chest X-ray Diagnosis Using Public Datasets"
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    add_para(doc, "Short title: NIH ChestX-ray14 to VinDr-CXR public-subset validation")
    add_para(doc, "Study type: retrospective public-dataset machine-learning validation study")

    doc.add_heading("Abstract", level=1)
    abstract_parts = [
        f"Background: Chest radiograph classifiers often show dataset-dependent behavior, and discrimination alone is insufficient for evaluating clinical translation. This study examined cross-dataset generalizability, probability calibration, and failure-mode saliency for public chest radiograph diagnosis.",
        f"Methods: A reproducible PyTorch pipeline was built using an executable public NIH ChestX-ray14 subset as the development dataset and a VinDr-CXR/VinBigData-derived public PNG subset as independent external validation. The NIH subset contained {reg['dataset']['nih_images']} images from {reg['dataset']['nih_patients']} patients and was split at patient level into training, validation, calibration, and internal-test partitions. The external subset contained {reg['dataset']['external_images']} images. DenseNet121 was prespecified as the main model; ResNet50 and EfficientNet-B0 were trained as architecture comparators. Thresholds were selected only on the NIH validation split, and calibration models were fitted only on the NIH calibration split.",
        f"Results: DenseNet121 achieved internal macro AUROC/AUPRC of {reg['performance']['internal_macro_auroc']}/{reg['performance']['internal_macro_auprc']} and external macro AUROC/AUPRC of {reg['performance']['external_macro_auroc']}/{reg['performance']['external_macro_auprc']}. Architecture-comparator results were: {comparator_sentence}. DenseNet121 internal uncalibrated macro Brier score, ECE, and MCE were {reg['calibration']['internal_uncalibrated_macro_brier']}, {reg['calibration']['internal_uncalibrated_macro_ece']}, and {reg['calibration']['internal_uncalibrated_macro_mce']}, respectively. Grad-CAM++ generated {reg['gradcam']['n_cases']} TP/FP/FN/TN examples for failure-mode review.",
        "Conclusions: The experiment provides a reproducible public-dataset validation package for examining discrimination, calibration, and saliency behavior under dataset shift. The findings support cautious model-comparison and calibration-transfer reporting, but they do not establish clinical readiness or full-cohort performance.",
    ]
    for part in abstract_parts:
        add_para(doc, part)
    add_para(doc, "Keywords: chest radiography; external validation; calibration; Grad-CAM; NIH ChestX-ray14; VinDr-CXR; DenseNet; ResNet")

    add_section_paragraphs(
        doc,
        "Introduction",
        [
            "Deep-learning systems for chest radiograph classification can achieve useful discrimination within a development dataset, but performance may change when applied to images from different institutions, labeling processes, or acquisition pipelines. For public chest radiograph benchmarks such as NIH ChestX-ray14 and VinDr-CXR, this problem is particularly relevant because labels may be derived from reports in one dataset and from radiologist annotations in another [1,2].",
            "External validation is therefore central to evaluating whether a model has learned disease-relevant patterns or dataset-specific correlates. Calibration adds a second requirement: a model with acceptable ranking performance may still produce probabilities that are poorly aligned with observed frequencies. Explainability methods such as Grad-CAM can help inspect failure modes, but saliency maps should not be interpreted as proof of clinical reasoning.",
            "This study developed an end-to-end, reproducible public-dataset pipeline to evaluate generalizability, calibration, and Grad-CAM++ failure modes for multi-label chest radiograph classification. DenseNet121 was used as the prespecified main model, with ResNet50 and EfficientNet-B0 as comparators after the main analysis was operational.",
        ],
    )

    add_section_paragraphs(
        doc,
        "Methods",
        [
            f"The development dataset was an NIH ChestX-ray14 public parquet subset extracted to PNG images [1]. The final development metadata contained {reg['dataset']['nih_images']} images from {reg['dataset']['nih_patients']} patients. The independent external validation dataset was a public VinDr-CXR/VinBigData-derived PNG subset containing {reg['dataset']['external_images']} images with image-level patient identifiers [2].",
            "The target labels were Atelectasis, Cardiomegaly, Pleural Effusion, Pneumothorax, Consolidation, Edema, Pneumonia, and No Finding. Edema and Pneumonia were retained for internal sensitivity analyses, but external primary claims were restricted when VinDr-derived labels were unavailable or had zero positive prevalence.",
            "NIH data were split at patient level into training, validation, calibration, and internal-test splits. The validation split was used for threshold selection and model selection. The calibration split was used for temperature scaling, Platt scaling, and isotonic regression. The internal-test split and external VinDr-derived subset were not used for training, threshold tuning, or calibration fitting.",
            "Images were resized to 224 x 224 pixels and normalized with ImageNet statistics. The training pipeline used BCEWithLogitsLoss with positive class weights, AdamW optimization, mixed precision on CUDA, and validation macro AUPRC for checkpoint selection. DenseNet121 was the main model; ResNet50 and EfficientNet-B0 were trained under the same split and evaluation rules as architecture comparators [3,4,7].",
            "Evaluation metrics included AUROC, AUPRC, sensitivity, specificity, F1-score, accuracy, positive predictive value, and negative predictive value, with macro and micro averages. Patient-level bootstrap resampling was used to estimate 95% confidence intervals. Calibration was assessed using Brier score, expected calibration error, maximum calibration error, calibration slope, and calibration intercept [5]. Grad-CAM++ was used only for qualitative failure-mode analysis [6].",
        ],
    )

    doc.add_heading("Results", level=1)
    add_para(
        doc,
        f"The NIH patient-level split produced {table1.loc[table1['split']=='train','n_images'].iloc[0]} training images, {table1.loc[table1['split']=='validation','n_images'].iloc[0]} validation images, {table1.loc[table1['split']=='calibration','n_images'].iloc[0]} calibration images, and {table1.loc[table1['split']=='internal_test','n_images'].iloc[0]} internal-test images. No patient overlap was detected among NIH splits.",
    )
    add_para(
        doc,
        f"DenseNet121 internal macro AUROC/AUPRC were {reg['performance']['internal_macro_auroc']}/{reg['performance']['internal_macro_auprc']}; external macro AUROC/AUPRC were {reg['performance']['external_macro_auroc']}/{reg['performance']['external_macro_auprc']}. The internal-minus-external macro AUROC difference was {reg['performance']['macro_auroc_absolute_drop']}, indicating no observed AUROC degradation in this selected external subset.",
    )
    add_para(
        doc,
        f"Architecture comparators produced the following macro discrimination results: {comparator_sentence}.",
    )
    add_para(
        doc,
        f"Before calibration, DenseNet121 internal macro Brier score was {reg['calibration']['internal_uncalibrated_macro_brier']}, ECE was {reg['calibration']['internal_uncalibrated_macro_ece']}, and MCE was {reg['calibration']['internal_uncalibrated_macro_mce']}. Temperature scaling yielded internal macro ECE {reg['calibration']['internal_temperature_macro_ece']} and external macro ECE {reg['calibration']['external_temperature_macro_ece']}. Calibration findings were interpreted with slope and intercept because bin-wise ECE alone can be misleading for sparse labels.",
    )
    add_para(
        doc,
        f"Grad-CAM++ panels were generated for {reg['gradcam']['n_cases']} true-positive, false-positive, false-negative, and true-negative cases across internal and external validation. The panels were used to review failure modes rather than to assert clinically valid causal reasoning.",
    )

    add_section_paragraphs(
        doc,
        "Discussion",
        [
            "This public-dataset experiment demonstrates a complete and auditable workflow for assessing discrimination, calibration, and saliency behavior in chest radiograph classification. The external subset showed higher macro AUPRC than the internal NIH test set, which should not be interpreted as superior generalization in a universal sense; it likely reflects differences in case mix, prevalence, label definitions, and selected public-file availability.",
            "Calibration results were mixed. Temperature scaling, Platt scaling, and isotonic regression were fitted only on the NIH calibration split and then transferred to internal and external test data. Improvements in ECE for some methods must be weighed against slope and intercept instability, particularly for labels with low prevalence or compressed probability ranges.",
            "The architecture comparators supported the main-model interpretation by showing the same broad pattern of external transfer in this selected public subset, but DenseNet121 remained the prespecified model. This design reduces single-architecture dependence without changing the main analysis after observing external performance.",
            "Grad-CAM++ visualizations were informative for inspecting apparent focus patterns in false positives and false negatives, but such heatmaps are coarse, model-dependent, and sensitive to preprocessing. They should be treated as exploratory failure-mode artifacts, not as evidence that the model used the same reasoning as a radiologist.",
        ],
    )

    add_section_paragraphs(
        doc,
        "Limitations",
        [
            "This is an executable public-subset analysis, not a full official-cohort analysis of all NIH ChestX-ray14 and VinDr-CXR images. The external data were derived from available public PNG files, and label availability differed by dataset.",
            "NIH ChestX-ray14 labels are report-mined weak labels, whereas VinDr-CXR annotations are radiologist-derived. This label-source heterogeneity can affect both discrimination and calibration estimates.",
            "Edema and Pneumonia had no positive cases in the external metadata used here and should not support external primary claims. These labels are retained only to preserve the internal sensitivity-label workflow.",
            "The models used 224 x 224 images and short training schedules suitable for a reproducible public experiment. Results should not be presented as the ceiling performance of these architectures.",
        ],
    )

    add_section_paragraphs(
        doc,
        "Conclusion",
        [
            "A reproducible NIH ChestX-ray14 to VinDr-CXR public-subset pipeline was completed for multi-label chest radiograph classification. DenseNet121, ResNet50, and EfficientNet-B0 were evaluated with validation-derived thresholds, patient-level bootstrap confidence intervals, NIH-only calibration fitting, and Grad-CAM++ failure-mode review. The study supports transparent public-dataset reporting of generalizability and calibration, while avoiding clinical-readiness claims.",
        ],
    )

    add_section_paragraphs(
        doc,
        "Data and Code Availability",
        [
            "All generated metadata, split files, prediction CSVs, metric tables, figures, logs, configuration files, and model checkpoints are packaged in reproducibility_package_hf_large.zip. The package also contains the random seed, GPU/environment reports, and rerun commands. The source data are public NIH ChestX-ray14 and VinDr-CXR/VinBigData-derived public files; users should follow the original dataset license and citation requirements.",
        ],
    )

    add_section_paragraphs(
        doc,
        "Ethics Statement",
        [
            "This analysis used de-identified public datasets and did not involve new patient recruitment or direct patient contact. Any institutional submission should confirm whether local review-board exemption or waiver documentation is required.",
        ],
    )

    doc.add_heading("Tables", level=1)
    add_df(doc, "Table 1. Dataset Characteristics", table1, ["split", "n_patients", "n_images"])
    add_df(doc, "Table 2. Label Harmonization", label_map, ["harmonized_label", "nih_terms", "vindr_terms", "role"], max_rows=8)
    add_df(doc, "Table 3. Internal and External Performance by Label and Model", compact_table3(t3), ["model", "dataset", "label", "prevalence", "AUROC (95% CI)", "AUPRC (95% CI)", "F1"], max_rows=48)
    add_df(doc, "Table 4. Macro-level Model Comparison", macro_display(macro), ["model", "internal AUROC", "internal AUPRC", "external AUROC", "external AUPRC", "AUROC drop", "AUPRC drop"])
    table5 = subgroup[["subgroup_variable", "subgroup", "n_images", "n_patients", "macro_auroc", "macro_auprc", "macro_f1", "macro_accuracy"]]
    add_df(doc, "Table 5. Sensitivity and Subgroup Analysis", table5, list(table5.columns))

    doc.add_heading("Figures", level=1)
    add_figure(doc, "Figure 1. Study Workflow", figs / "figure1_study_workflow_large.png", "Patient-level NIH development workflow with independent VinDr-derived external validation.")
    add_figure(doc, "Figure 2. Internal vs External AUROC/AUPRC", figs / "figure2_internal_vs_external_auroc_auprc_ci.png", "Label-wise DenseNet121 internal and external discrimination metrics.")
    add_figure(doc, "Figure 3. Calibration Curves", figs / "figure3_calibration_curves_large.png", "Reliability diagrams after NIH-only calibration fitting and transfer.")
    add_figure(doc, "Figure 4. Grad-CAM++ Failure-mode Examples", figs / "figure4_gradcam_failure_modes_large.png", "Representative TP/FP/FN/TN heatmap panels for qualitative failure-mode analysis.")
    add_figure(doc, "Figure 5. External Performance Drop by Label", figs / "figure5_external_performance_drop_by_label.png", "Internal minus external AUROC; negative values indicate higher AUROC in the selected external subset.")

    doc.add_heading("References", level=1)
    refs = [
        "1. Wang X, Peng Y, Lu L, Lu Z, Bagheri M, Summers RM. ChestX-ray8: Hospital-scale chest X-ray database and benchmarks on weakly-supervised classification and localization of common thorax diseases. CVPR. 2017.",
        "2. Nguyen HQ, Lam K, Le LT, et al. VinDr-CXR: An open dataset of chest X-rays with radiologist's annotations. Scientific Data. 2022;9:429.",
        "3. Huang G, Liu Z, van der Maaten L, Weinberger KQ. Densely connected convolutional networks. CVPR. 2017.",
        "4. He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition. CVPR. 2016.",
        "5. Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. ICML. 2017.",
        "6. Selvaraju RR, Cogswell M, Das A, Vedantam R, Parikh D, Batra D. Grad-CAM: Visual explanations from deep networks via gradient-based localization. ICCV. 2017.",
        "7. Tan M, Le QV. EfficientNet: Rethinking model scaling for convolutional neural networks. ICML. 2019.",
    ]
    for ref in refs:
        add_para(doc, ref)

    doc.save(path)
    return path


if __name__ == "__main__":
    out = make_manuscript(DESKTOP / "CXR_Manuscript_Version_C_RouteB_Large_SCI_Submission.docx")
    shutil.copy2(out, ROOT / "outputs_large" / out.name)
    print(out)


