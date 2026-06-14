from __future__ import annotations

import json
import math
import shutil
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
TABLES = ROOT / "tables_large"
FIGS = ROOT / "figures_large"
OUTPUTS = ROOT / "outputs_large"

LABELS = [
    "Atelectasis",
    "Cardiomegaly",
    "Pleural Effusion",
    "Pneumothorax",
    "Consolidation",
    "Edema",
    "Pneumonia",
    "No Finding",
]


def fmt(x: object, digits: int = 3) -> str:
    try:
        val = float(x)
    except Exception:
        return str(x)
    if not math.isfinite(val):
        return "NA"
    return f"{val:.{digits}f}"


def table(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


def registry() -> dict:
    return json.loads((OUTPUTS / "manuscript_value_registry_10_upgrade.json").read_text(encoding="utf-8"))


def env_report() -> dict:
    return json.loads((OUTPUTS / "training" / "hf_large_densenet121" / "environment_report.json").read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para(doc, item, style="List Bullet")


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(15)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(10)


def style_table(table_obj) -> None:
    table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_obj.style = "Table Grid"
    table_obj.autofit = True
    for i, row in enumerate(table_obj.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if i == 0:
                set_cell_shading(cell, "F2F4F7")
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(7 if len(table_obj.columns) >= 8 else 8)
                    if i == 0:
                        run.bold = True


def add_df(doc: Document, title: str, df: pd.DataFrame, columns: list[str], max_rows: int | None = None) -> None:
    doc.add_heading(title, level=2)
    use = df[columns].copy()
    if max_rows is not None:
        use = use.head(max_rows)
    tbl = doc.add_table(rows=1, cols=len(columns))
    for idx, col in enumerate(columns):
        tbl.rows[0].cells[idx].text = col
    for _, row in use.iterrows():
        cells = tbl.add_row().cells
        for idx, col in enumerate(columns):
            value = row[col]
            cells[idx].text = fmt(value) if isinstance(value, float) else str(value)
    style_table(tbl)


def image_for_docx(path: Path) -> Path:
    from PIL import Image

    out_dir = OUTPUTS / "docx_embedded_images_12_final"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{path.stem}.jpg"
    if out.exists() and out.stat().st_mtime >= path.stat().st_mtime:
        return out
    im = Image.open(path).convert("RGB")
    im.save(out, quality=94, optimize=True)
    return out


def add_figure(doc: Document, title: str, path: Path, caption: str, width: float = 6.35) -> None:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_for_docx(path)), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.style = "Caption"


def build_final12_tables() -> dict[str, pd.DataFrame]:
    reg = registry()
    t1_old = table("table1_dataset_characteristics_upgraded.csv")
    t2_old = table("table2_label_harmonization_analysis_role_upgraded.csv")
    t3_old = table("table3_densenet121_labelwise_performance.csv")
    t4_old = table("table4_architecture_comparison_macro_upgraded.csv")
    t5_old = table("table5_calibration_metrics_upgraded.csv")
    t6_old = table("table6_subgroup_sensitivity_analysis.csv")

    rows = []
    for _, row in t1_old.iterrows():
        split = row["split"]
        role = row["role"]
        if role == "development":
            split_level = "patient-level" if split != "overall" else "patient-level split applied within NIH"
            ident = row["n_patients"]
        else:
            split_level = "image-identifier-level external subset"
            ident = row["n_patients"]
        rows.append(
            {
                "dataset": row["dataset"],
                "role": role,
                "split": split,
                "n_images": int(row["n_images"]),
                "n_patients_or_identifiers": int(ident),
                "split_level": split_level,
                "used_for_training": "yes" if split == "train" else "no",
                "used_for_threshold_selection": "yes" if split == "validation" else "no",
                "used_for_calibration_fitting": "yes" if split == "calibration" else "no",
                "used_for_final_evaluation": "yes" if split in {"internal_test", "external"} else "no",
            }
        )
    t1 = pd.DataFrame(rows)

    t2 = pd.DataFrame(
        {
            "harmonized_label": t2_old["harmonized_label"],
            "NIH_terms": t2_old["nih_label_source"],
            "VinDr_terms": t2_old["vindr_label_source"],
            "internal_analysis_role": t2_old["analysis_role"].str.replace("; external-evaluable", "", regex=False),
            "external_positive_cases": t2_old["external_positives_in_analyzed_subset"],
            "external_evaluable": t2_old["external_evaluable_for_auc_pr"],
            "external_claim_status": t2_old["external_evaluable_for_auc_pr"].map(lambda x: "external primary/sensitivity claim allowed" if bool(x) else "not used for external discrimination claims"),
            "caution": t2_old["caution"].fillna(""),
        }
    )

    t3 = pd.DataFrame(
        {
            "label": t3_old["label"],
            "validation_threshold": t3_old["validation_threshold"],
            "internal_prevalence": t3_old["nih_internal_prevalence"],
            "external_prevalence": t3_old["vindr_external_prevalence"],
            "internal_AUROC_95CI": t3_old["internal_auroc_95ci"],
            "external_AUROC_95CI": t3_old["external_auroc_95ci"],
            "internal_AUPRC_95CI": t3_old["internal_auprc_95ci"],
            "external_AUPRC_95CI": t3_old["external_auprc_95ci"],
            "external_status": t3_old["external_analysis_status"],
        }
    )

    t4 = pd.DataFrame(
        {
            "model": t4_old["model"],
            "internal_all_label_macro_AUROC": t4_old["internal_all_labels_macro_auroc"],
            "internal_all_label_macro_AUPRC": t4_old["internal_all_labels_macro_auprc"],
            "internal_external_evaluable_macro_AUROC": t4_old["internal_external_evaluable_labels_macro_auroc"],
            "internal_external_evaluable_macro_AUPRC": t4_old["internal_external_evaluable_labels_macro_auprc"],
            "external_nonzero_prevalence_macro_AUROC": t4_old["external_nonzero_prevalence_labels_macro_auroc"],
            "external_nonzero_prevalence_macro_AUPRC": t4_old["external_nonzero_prevalence_labels_macro_auprc"],
        }
    )

    t5 = pd.DataFrame(
        {
            "dataset": t5_old["dataset"],
            "calibration_method": t5_old["calibration_method"],
            "Brier_score": t5_old["brier_score"],
            "ECE": t5_old["ece"],
            "MCE": t5_old["mce"],
            "calibration_slope": t5_old["calibration_slope"],
            "calibration_intercept": t5_old["calibration_intercept"],
            "external_fitting_used": t5_old["external_fitting_used"],
            "interpretation": "NIH-fitted method evaluated for transfer; interpret with Brier, ECE, MCE, slope, intercept, and prevalence",
        }
    )

    t6 = pd.DataFrame(
        {
            "dataset": t6_old["dataset"],
            "subgroup_variable": t6_old["subgroup_variable"],
            "subgroup": t6_old["subgroup"],
            "n_images": t6_old["n_images"],
            "n_patients_or_identifiers": t6_old["n_patients"],
            "macro_AUROC": t6_old["all_label_macro_auroc"],
            "macro_AUPRC": t6_old["all_label_macro_auprc"],
            "analysis_status": t6_old["analysis_status"],
        }
    )

    out = {
        "table1_dataset_characteristics_final12.csv": t1,
        "table2_label_harmonization_final12.csv": t2,
        "table3_densenet121_labelwise_performance_final12.csv": t3,
        "table4_architecture_comparison_final12.csv": t4,
        "table5_calibration_metrics_final12.csv": t5,
        "table6_subgroup_sensitivity_final12.csv": t6,
    }
    for name, df in out.items():
        df.to_csv(TABLES / name, index=False)
    return out


def title_options() -> list[dict[str, str]]:
    return [
        {
            "title": "Public-subset Cross-dataset Validation and Calibration Transfer of Deep Learning Models for Multi-label Chest Radiograph Classification",
            "comment": "Recommended. It is concise, states public-subset scope, and includes calibration transfer without implying clinical deployment.",
        },
        {
            "title": "Cross-dataset Discrimination, Calibration Transfer, and Failure-mode Saliency in Public-subset Chest Radiograph Classification",
            "comment": "Balanced option with strong emphasis on the three central analyses.",
        },
        {
            "title": "From NIH ChestX-ray14 to a VinDr-derived External Subset: Public-subset Validation of Chest Radiograph Classifiers",
            "comment": "Dataset-forward option, useful for applied medical imaging journals.",
        },
        {
            "title": "Calibration Transfer and Grad-CAM++ Failure-mode Review in Public-subset Chest Radiograph AI Validation",
            "comment": "Highlights calibration and saliency but is less explicit about architecture comparison.",
        },
        {
            "title": "Generalizability and Calibration of Multi-label Chest Radiograph Classifiers in Selected Public NIH and VinDr-derived Subsets",
            "comment": "Short and cautious, though less specific about Grad-CAM++.",
        },
    ]


def manuscript_sections() -> dict[str, list[str] | str]:
    reg = registry()
    env = env_report()
    d = reg["densenet121"]
    t4 = table("table4_architecture_comparison_macro_upgraded.csv")
    paired = table("supplementary_table6_paired_bootstrap_model_comparison.csv")
    paired_internal = paired[
        (paired["dataset"].eq("NIH internal test"))
        & (paired["label_set"].eq("all_valid_labels"))
        & (paired["metric"].eq("auroc"))
        & (paired["model_a"].eq("DenseNet121"))
        & (paired["model_b"].eq("ResNet50"))
    ].iloc[0]

    title = title_options()[0]["title"]
    short_title = "Public-subset chest radiograph validation"
    abstract = [
        "Background: Chest radiograph classifiers are vulnerable to dataset shift, case-mix differences, label-source heterogeneity, and imperfect probability calibration. Cross-dataset validation should therefore evaluate discrimination, prevalence-aware precision-recall behavior, calibration transfer, and qualitative failure modes without implying clinical readiness.",
        f"Methods: We conducted a retrospective public-dataset machine-learning validation study using a selected NIH ChestX-ray14 public parquet subset as the development dataset and a VinDr-CXR/VinBigData-derived public PNG subset as independent external validation. The NIH subset contained {reg['dataset']['nih_images']} images from {reg['dataset']['nih_patients']} patients and was split at patient level into training, validation, calibration, and internal-test partitions. The external subset contained {reg['dataset']['vindr_images']} images. DenseNet121 was the prespecified main model, and ResNet50 and EfficientNet-B0 were architecture comparators. Per-label thresholds were selected only by maximizing F1 on the NIH validation split, and temperature scaling, Platt scaling, and isotonic regression were fitted only on the NIH calibration split.",
        f"Results: DenseNet121 achieved internal all-label macro AUROC/AUPRC of {d['internal_all_label_macro_auroc']}/{d['internal_all_label_macro_auprc']}. On the external-evaluable label set, NIH internal macro AUROC/AUPRC were {d['internal_external_evaluable_macro_auroc']}/{d['internal_external_evaluable_macro_auprc']}, whereas the external nonzero-prevalence macro AUROC/AUPRC were {d['external_nonzero_prevalence_macro_auroc']}/{d['external_nonzero_prevalence_macro_auprc']}. ResNet50 and EfficientNet-B0 external nonzero-prevalence macro AUROC/AUPRC were {fmt(t4.loc[t4['model'].eq('ResNet50'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}/{fmt(t4.loc[t4['model'].eq('ResNet50'), 'external_nonzero_prevalence_labels_macro_auprc'].iloc[0])} and {fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}/{fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'external_nonzero_prevalence_labels_macro_auprc'].iloc[0])}, respectively. Edema and Pneumonia had zero external positives and were not used for external discrimination claims. DenseNet121 internal uncalibrated macro Brier score, ECE, MCE, slope, and intercept were {reg['calibration']['internal_uncalibrated_brier']}, {reg['calibration']['internal_uncalibrated_ece']}, {reg['calibration']['internal_uncalibrated_mce']}, {reg['calibration']['internal_uncalibrated_slope']}, and {reg['calibration']['internal_uncalibrated_intercept']}. Calibration effects were method-dependent. Grad-CAM++ was used for qualitative failure-mode review of true-positive, false-positive, false-negative, and true-negative examples.",
        "Conclusions: This study provides an auditable public-subset framework for evaluating cross-dataset discrimination, calibration transfer, architecture comparators, and failure-mode saliency in multi-label chest radiograph classification. The results do not establish full official-cohort performance or clinical readiness.",
    ]
    introduction = [
        "Chest radiography is among the most frequently used imaging examinations, and public chest radiograph datasets have enabled reproducible research on multi-label deep learning models. However, performance measured within a development dataset may not translate directly to images collected under different acquisition, annotation, and selection processes.",
        "Public chest radiograph datasets differ in domain characteristics, disease prevalence, case mix, and label generation. NIH ChestX-ray14 uses report-mined weak labels, whereas VinDr-CXR provides radiologist-derived annotations. These differences are scientifically useful for external validation but create label-source heterogeneity that must be reflected in the interpretation.",
        "Discrimination metrics alone are incomplete. AUROC evaluates ranking across thresholds, whereas AUPRC is strongly affected by positive-label prevalence and is often more informative for imbalanced findings. Thresholded metrics depend on where thresholds are selected, and predicted probabilities require calibration assessment before they can be interpreted as risk-like outputs.",
        "Saliency methods such as Grad-CAM++ can support qualitative inspection of model failures, including false positives and false negatives. These maps are coarse and method-dependent, and they should not be interpreted as localization ground truth or proof that a model reasons like a radiologist.",
        "We therefore developed a reproducible public-subset pipeline to evaluate cross-dataset discrimination, prevalence-aware precision-recall behavior, calibration transfer, architecture comparators, and Grad-CAM++ failure modes for multi-label chest radiograph classification from NIH ChestX-ray14 to a selected VinDr-derived external subset.",
    ]
    methods = [
        ("Study design", "This was a retrospective public-dataset machine-learning validation study. All reported numeric results were derived from saved prediction CSV files and metric tables generated by the reproducible project pipeline."),
        ("Data sources and subset construction", f"The development dataset was a selected NIH ChestX-ray14 public parquet subset containing {reg['dataset']['nih_images']} images from {reg['dataset']['nih_patients']} patients. The independent external validation dataset was a VinDr-CXR/VinBigData-derived public PNG subset containing {reg['dataset']['vindr_images']} images. The external subset was treated as a selected public external subset rather than the full official VinDr-CXR cohort."),
        ("Label harmonization and externally evaluable labels", "The harmonized label set comprised Atelectasis, Cardiomegaly, Pleural Effusion, Pneumothorax, Consolidation, Edema, Pneumonia, and No Finding. External primary discrimination claims were restricted to external-evaluable labels with both positive and negative cases in the selected external subset: Atelectasis, Cardiomegaly, Pleural Effusion, Pneumothorax, Consolidation, and No Finding. Edema and Pneumonia had zero external positives, so their external AUROC and AUPRC were not calculable and were reported as not applicable."),
        ("Data splitting and leakage control", "NIH data were split at patient level into training, validation, calibration, and internal-test partitions. The validation split was used for checkpoint selection and per-label threshold selection. The calibration split was used only for fitting calibration models. Internal-test and external data were held out from training, threshold tuning, checkpoint selection, and calibration fitting."),
        ("Image preprocessing and augmentation", "Images were resized to 224 x 224 pixels, converted to three-channel tensors, and normalized with ImageNet mean and standard deviation. Training augmentation included horizontal flipping with probability 0.5 and random rotation within 7 degrees. The same deterministic resizing and normalization were used for evaluation."),
        ("Model architectures and training", f"DenseNet121 was the prespecified main architecture, and ResNet50 and EfficientNet-B0 were architecture comparators. All models used ImageNet-pretrained torchvision weights. Training used BCEWithLogitsLoss with positive class weights calculated as negative cases divided by positive cases for each label, clamped to the range 1 to 100. AdamW was used with learning rate 0.0001 and weight decay 0.0001. Training used batch size 24, evaluation batch size 48, gradient clipping at norm 1.0, mixed precision when CUDA was available, and early stopping patience 3 based on validation macro AUPRC. The scheduler was ReduceLROnPlateau with mode=max, factor 0.5, and patience 2. The run used 5 planned epochs, random seed 20260614, Python {env.get('python', '').split()[0]}, PyTorch {env.get('torch', '')}, CUDA available, and GPU {env.get('gpu_name', '')}."),
        ("Threshold selection", "Per-label thresholds were selected by maximizing F1 on the NIH validation split. Thresholds were not re-optimized on the NIH internal-test split or the VinDr-derived external subset."),
        ("Calibration methods", "Temperature scaling, Platt scaling, and isotonic regression were fitted only on the NIH calibration split. No external calibration fitting was performed in the primary analysis. Calibration transfer was evaluated on NIH internal-test and VinDr external predictions."),
        ("Evaluation metrics", "Discrimination metrics included AUROC and AUPRC. Thresholded metrics included sensitivity, specificity, F1-score, accuracy, positive predictive value, and negative predictive value. Calibration metrics included Brier score, expected calibration error, maximum calibration error, calibration slope, and calibration intercept. Macro averages excluded labels for which AUROC or AUPRC were not calculable. AUPRC baseline and lift used observed prevalence, with lift defined as AUPRC divided by prevalence."),
        ("Statistical analysis", "Bootstrap confidence intervals used 500 iterations and percentile intervals. Bootstrap confidence intervals were estimated at the patient level for NIH internal testing. For the cached external subset, where reliable patient-level identifiers were not available, resampling was performed at the image-identifier level and interpreted as an approximation. Paired bootstrap model comparisons were descriptive rather than confirmatory superiority tests."),
        ("Grad-CAM++ failure-mode review", "Grad-CAM++ heatmaps were generated for representative true-positive, false-positive, false-negative, and true-negative predictions in NIH internal and VinDr external evaluation. These saliency maps were used only for qualitative failure-mode inspection and not as localization ground truth or evidence of radiologist-like reasoning."),
        ("Reproducibility package", "The reproducibility package includes split CSVs, label mapping, prediction CSVs, calibration outputs, figure source data, model checkpoints, training logs, package versions, random seeds, GPU/environment information, and rerun commands."),
        ("Ethics statement", "The study used de-identified public datasets, involved no new patient recruitment, no intervention, and no direct patient contact. Local institutional review-board exemption or waiver requirements should be confirmed before journal submission."),
    ]
    results = [
        ("Dataset characteristics", "The NIH development subset contained 5000 images from 3795 patients. The patient-level split yielded 2679 training images, 831 validation images, 751 calibration images, and 739 internal-test images. The selected VinDr-derived external subset contained 1000 images."),
        ("Externally evaluable labels", "The external-evaluable labels were Atelectasis, Cardiomegaly, Pleural Effusion, Pneumothorax, Consolidation, and No Finding. Edema and Pneumonia had zero positive cases in the selected external subset; their external AUROC and AUPRC were therefore not calculable and were not used for external primary discrimination claims."),
        ("DenseNet121 discrimination performance", f"DenseNet121 achieved internal all-label macro AUROC/AUPRC of {d['internal_all_label_macro_auroc']}/{d['internal_all_label_macro_auprc']}. On the same external-evaluable label set used for external macro summaries, NIH internal macro AUROC/AUPRC were {d['internal_external_evaluable_macro_auroc']}/{d['internal_external_evaluable_macro_auprc']}. External nonzero-prevalence macro AUROC/AUPRC were {d['external_nonzero_prevalence_macro_auroc']}/{d['external_nonzero_prevalence_macro_auprc']}. The external AUPRC should be interpreted with the external prevalence distribution and selected subset composition rather than as evidence of stronger external generalization."),
        ("Architecture comparison", f"ResNet50 achieved internal all-label macro AUROC/AUPRC of {fmt(t4.loc[t4['model'].eq('ResNet50'), 'internal_all_labels_macro_auroc'].iloc[0])}/{fmt(t4.loc[t4['model'].eq('ResNet50'), 'internal_all_labels_macro_auprc'].iloc[0])} and external nonzero-prevalence macro AUROC/AUPRC of {fmt(t4.loc[t4['model'].eq('ResNet50'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}/{fmt(t4.loc[t4['model'].eq('ResNet50'), 'external_nonzero_prevalence_labels_macro_auprc'].iloc[0])}. EfficientNet-B0 achieved internal all-label macro AUROC/AUPRC of {fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'internal_all_labels_macro_auroc'].iloc[0])}/{fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'internal_all_labels_macro_auprc'].iloc[0])} and external nonzero-prevalence macro AUROC/AUPRC of {fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}/{fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'external_nonzero_prevalence_labels_macro_auprc'].iloc[0])}. The paired bootstrap comparison favored DenseNet121 over ResNet50 for internal all-label AUROC only (difference {fmt(paired_internal['difference_a_minus_b'])}, 95% CI {fmt(paired_internal['ci_lower'])} to {fmt(paired_internal['ci_upper'])}); most external-evaluable AUROC and AUPRC differences crossed zero."),
        ("Prevalence-aware AUPRC interpretation", "AUPRC baseline and lift analyses were used to contextualize precision-recall behavior under class imbalance. Because AUPRC is prevalence-sensitive, the numerically higher external AUPRC should be interpreted as a selected-subset result affected by prevalence, case mix, and label-source differences."),
        ("Calibration transfer", "Calibration effects were method-dependent. On NIH internal testing, the uncalibrated DenseNet121 macro Brier score, ECE, MCE, calibration slope, and intercept were 0.193, 0.269, 0.539, 0.633, and -2.285, respectively. Temperature scaling did not consistently reduce ECE. Platt scaling and isotonic regression reduced apparent Brier score and ECE in this experiment, but slope and intercept behavior indicated that probability-scale interpretation should remain cautious, especially under external transfer."),
        ("Subgroup and sensitivity analysis", "Available NIH sex and view-position subgroup summaries were descriptive and not powered for fairness inference. The cached external subset lacked reliable demographic metadata for comparable external subgroup analysis."),
        ("Grad-CAM++ failure-mode review", "Grad-CAM++ panels were generated for true-positive, false-positive, false-negative, and true-negative examples in internal and external evaluation. The panels supported qualitative failure-mode inspection but did not provide localization ground truth or proof of radiologist-like model reasoning."),
    ]
    discussion = [
        "This study completed a reproducible public-subset validation pipeline for multi-label chest radiograph classification, including DenseNet121 as the prespecified main model, ResNet50 and EfficientNet-B0 architecture comparators, NIH-only calibration fitting, external validation, and Grad-CAM++ failure-mode review.",
        "The internal and external differences should be interpreted cautiously. The numerically higher external macro AUROC/AUPRC in this selected public external subset may reflect selected subset composition, positive-label prevalence, case mix, and label-source heterogeneity rather than superior cross-dataset generalization.",
        "The AUPRC baseline and lift analyses are important because precision-recall summaries are prevalence-dependent. A higher external AUPRC can occur when the external subset contains a different distribution of positive cases, and it should not be interpreted without the observed prevalence baseline.",
        "The architecture comparison reduces dependence on a single model family but does not establish a definitive architecture ranking. DenseNet121 remained the prespecified main model, whereas ResNet50 and EfficientNet-B0 provided descriptive comparator evidence under the same split and evaluation rules.",
        "Calibration transfer remained method-dependent. Temperature scaling was not consistently helpful in this experiment. Platt scaling and isotonic regression reduced some apparent Brier and ECE values, but calibration slope and intercept suggested that probability-scale behavior was not fully resolved, especially when transferred to the selected external subset.",
        "Grad-CAM++ saliency panels were useful for illustrating failure modes, but they do not validate localization or clinical reasoning. Their appropriate role is qualitative inspection of representative errors rather than evidence that the model uses radiologist-like diagnostic logic.",
        "The practical contribution is an auditable workflow and reproducibility package for public-subset chest radiograph AI validation. The study supports transparent reporting of discrimination, calibration transfer, and failure-mode saliency, but it does not establish full official-cohort performance, prospective validity, or clinical readiness.",
    ]
    limitations = [
        "This analysis used selected public subsets rather than the full official NIH ChestX-ray14 and VinDr-CXR cohorts. The external subset may therefore introduce selection bias and should be described as a selected public external subset.",
        "The NIH and VinDr-derived labels were produced through different annotation processes: NIH labels are report-mined weak labels, whereas VinDr-CXR labels are radiologist-derived. This label-source heterogeneity can influence both discrimination and calibration estimates.",
        "Edema and Pneumonia had zero positive cases in the selected external subset, so external discrimination metrics for these labels were not calculable and should not be used as external primary claims.",
        "AUPRC is prevalence-sensitive, and external precision-recall performance should be interpreted with observed prevalence and AUPRC lift rather than compared directly with internal all-label macro AUPRC.",
        "Images were resized to 224 x 224 pixels, and the compact training schedule was designed for reproducible public-subset validation rather than leaderboard-optimized full-resolution performance.",
        "The study did not include prospective validation, reader study comparison, or clinical workflow evaluation. External patient identifiers were limited in the cached external subset, so external bootstrap estimates used image-identifier-level resampling as an approximation.",
        "Calibration transfer may be unstable across datasets, particularly for sparse labels and different annotation sources. Grad-CAM++ was qualitative only, and subgroup/fairness analysis was limited by available metadata.",
    ]
    conclusion = "A reproducible NIH ChestX-ray14 to VinDr-derived public-subset pipeline was completed for multi-label chest radiograph classification. DenseNet121, ResNet50, and EfficientNet-B0 were evaluated with validation-derived thresholds, patient-level NIH internal bootstrap intervals, image-identifier-level external bootstrap approximations, NIH-only calibration fitting, and Grad-CAM++ failure-mode review. The study supports transparent evaluation of cross-dataset discrimination, calibration transfer, and architecture comparison, but it does not establish full official-cohort performance or clinical readiness."
    data_code = "Split CSVs, label mapping, prediction CSVs, calibration outputs, figure source data, model checkpoints, logs, package versions, random seeds, GPU/environment information, and rerun commands are included in the reproducibility package. The public repository or archive link should be added before submission: [TO BE FILLED: repository or archive link]. Original images and labels remain under the source dataset licenses; users should follow NIH ChestX-ray14 and VinDr-CXR source license and citation requirements."
    ethics = "This study used de-identified public datasets and involved no new patient recruitment, no intervention, and no direct patient contact. Institutional review-board exemption or waiver requirements should be confirmed locally before journal submission."
    return {
        "title": title,
        "short_title": short_title,
        "abstract": abstract,
        "introduction": introduction,
        "methods": methods,
        "results": results,
        "discussion": discussion,
        "limitations": limitations,
        "conclusion": conclusion,
        "data_code": data_code,
        "ethics": ethics,
    }


def figure_captions() -> dict[str, str]:
    return {
        "Figure 1. Study workflow": "Study workflow for public-subset cross-dataset validation. A selected NIH ChestX-ray14 public parquet subset was used for development and split at patient level into training, validation, calibration, and internal-test partitions. The validation split was used for checkpoint and threshold selection, the NIH calibration split was used for post-hoc calibration fitting, and the internal-test split was used for held-out evaluation. A selected VinDr-CXR/VinBigData-derived public PNG subset was used only as independent external validation.",
        "Figure 2. DenseNet121 internal and external AUROC/AUPRC": "Label-wise DenseNet121 discrimination on NIH internal testing and the selected VinDr-derived external subset. Points and horizontal bars show point estimates and 95% bootstrap intervals. Edema and Pneumonia are marked not applicable externally because the analyzed external subset contained zero positive cases for these labels. Macro summaries were calculated separately for all internal labels and for the external-evaluable label set.",
        "Figure 3. AUPRC baseline and lift": "Prevalence-aware AUPRC interpretation. The prevalence baseline is the observed positive-label prevalence, and AUPRC lift is defined as AUPRC divided by prevalence. This panel contextualizes precision-recall behavior under class imbalance; external AUPRC should not be interpreted without considering the external prevalence distribution.",
        "Figure 4. Calibration curves": "Reliability curves for uncalibrated predictions and NIH-fitted temperature scaling, Platt scaling, and isotonic regression. Calibration methods were fitted only on the NIH calibration split and evaluated on NIH internal-test and selected VinDr external predictions. ECE should be interpreted together with Brier score, MCE, calibration slope, and calibration intercept.",
        "Figure 5. Grad-CAM++ failure-mode panels": "Representative Grad-CAM++ panels for true-positive, false-positive, false-negative, and true-negative predictions from internal and external evaluation. These heatmaps are qualitative failure-mode saliency examples only; they are not localization ground truth and do not demonstrate radiologist-like reasoning.",
        "Figure 6. Internal-minus-external AUROC difference": "Internal-minus-external AUROC differences for externally evaluable labels. Negative values indicate numerically higher AUROC in the selected external subset. Interpretation is affected by selected subset composition, prevalence, case mix, and label-source heterogeneity.",
    }


def table_plans() -> tuple[pd.DataFrame, pd.DataFrame]:
    main_rows = [
        ("Table 1. Dataset characteristics and split roles", "Document dataset role, split sizes, identifier level, and whether each split was used for training, threshold selection, calibration fitting, or final evaluation.", "dataset; role; split; n_images; n_patients_or_identifiers; split_level; used_for_training; used_for_threshold_selection; used_for_calibration_fitting; used_for_final_evaluation", "Main text", "Do not describe the VinDr-derived subset as full official cohort or patient-level if only image identifiers are reliable."),
        ("Table 2. Label harmonization and analysis role", "Make explicit which labels are externally evaluable and which labels cannot support external claims.", "harmonized_label; NIH_terms; VinDr_terms; internal_analysis_role; external_positive_cases; external_evaluable; external_claim_status; caution", "Main text", "Edema and Pneumonia must be marked zero-positive externally and not primary external labels."),
        ("Table 3. DenseNet121 label-wise internal and external performance", "Present main-model label-wise discrimination and thresholds.", "label; validation_threshold; internal_prevalence; external_prevalence; internal_AUROC_95CI; external_AUROC_95CI; internal_AUPRC_95CI; external_AUPRC_95CI; external_status", "Main text", "External NA values should remain NA; avoid replacing them with zero."),
        ("Table 4. Macro-level architecture comparison", "Compare DenseNet121 with ResNet50 and EfficientNet-B0 at macro level.", "model; internal_all_label_macro_AUROC; internal_all_label_macro_AUPRC; internal_external_evaluable_macro_AUROC; internal_external_evaluable_macro_AUPRC; external_nonzero_prevalence_macro_AUROC; external_nonzero_prevalence_macro_AUPRC", "Main text", "Describe model comparisons as descriptive unless paired bootstrap supports a narrow claim."),
        ("Table 5. Calibration metrics before and after NIH-fitted calibration", "Report calibration transfer across uncalibrated, temperature, Platt, and isotonic outputs.", "dataset; calibration_method; Brier_score; ECE; MCE; calibration_slope; calibration_intercept; external_fitting_used; interpretation", "Main text", "Do not say calibration improved globally; effects are method-dependent."),
        ("Table 6. Descriptive subgroup and sensitivity analysis", "Summarize available subgroup and sensitivity analyses.", "dataset; subgroup_variable; subgroup; n_images; n_patients_or_identifiers; macro_AUROC; macro_AUPRC; analysis_status", "Main text or supplement depending on journal space", "Do not infer fairness from underpowered descriptive subgroups."),
    ]
    supp_rows = [
        ("Supplementary Table 1. ResNet50 label-wise performance", "Comparator label-wise metrics.", "same structure as Table 3", "Supplement", "Comparator evidence only."),
        ("Supplementary Table 2. EfficientNet-B0 label-wise performance", "Comparator label-wise metrics.", "same structure as Table 3", "Supplement", "Comparator evidence only."),
        ("Supplementary Table 3. Per-label thresholds from NIH validation", "Record threshold provenance.", "label; threshold; validation_F1", "Supplement", "Thresholds were not optimized on internal-test or external data."),
        ("Supplementary Table 4. AUPRC prevalence baseline and lift", "Prevalence-aware interpretation of AUPRC.", "model; dataset; label; positives; prevalence; AUPRC; AUPRC_lift; status", "Supplement", "Use to avoid overinterpreting external AUPRC."),
        ("Supplementary Table 5. Per-label calibration metrics", "Detailed label-level calibration values.", "dataset; method; label; Brier; ECE; MCE; slope; intercept", "Supplement", "Sparse labels may have unstable slopes/intercepts."),
        ("Supplementary Table 6. Paired bootstrap model comparisons", "Descriptive architecture comparison.", "dataset; label_set; metric; model_a; model_b; difference; CI; n_bootstrap; interpretation", "Supplement", "Do not convert descriptive bootstrap into a superiority trial."),
        ("Supplementary Table 7. Hyperparameters and environment", "Reproducibility of training environment.", "parameter; value", "Supplement", "Must match logs/config."),
        ("Supplementary Table 8. Reproducibility commands", "Rerun commands.", "step; command", "Supplement", "Keep commands synchronized with repo."),
        ("Supplementary Table 9. Dataset-source and license summary", "Dataset provenance and use constraints.", "dataset; source; role; license/citation note", "Supplement", "Add official links and license text before journal submission."),
    ]
    cols = ["table_name", "purpose", "recommended_columns", "placement", "interpretation_risk"]
    return pd.DataFrame(main_rows, columns=cols), pd.DataFrame(supp_rows, columns=cols)


def overclaim_table() -> pd.DataFrame:
    rows = [
        ("proved", "Implies certainty beyond retrospective public-subset evidence.", "suggested within the evaluated subset"),
        ("robust", "Can imply broad transportability.", "consistent within this selected subset"),
        ("superior", "Implies confirmatory comparative testing.", "numerically higher or descriptively higher"),
        ("clinical utility", "Implies patient-care usefulness.", "relevance for model evaluation"),
        ("deployment-ready", "Implies clinical readiness.", "not evaluated for prospective clinical deployment"),
        ("radiologist-level", "Implies benchmark against clinical readers.", "not compared with radiologists in this study"),
        ("explainability proved reasoning", "Overstates saliency maps.", "saliency maps supported qualitative failure-mode inspection"),
        ("calibration improved reliability", "Suggests uniform improvement.", "calibration effects were method-dependent"),
        ("external generalization was better", "Ignores prevalence and selected subset composition.", "external metrics were numerically higher in this selected public subset"),
        ("full VinDr validation", "Misstates analyzed data scope.", "selected VinDr-derived public external subset"),
    ]
    return pd.DataFrame(rows, columns=["original_expression", "risk", "recommended_replacement"])


def checklist_table() -> pd.DataFrame:
    rows = []
    items = {
        "A. Scientific validity": [
            "State public-subset scope throughout.",
            "Avoid clinical-readiness claims.",
            "Keep DenseNet121 as prespecified main model.",
            "Describe ResNet50/EfficientNet-B0 as architecture comparators.",
            "Report external findings only for external-evaluable labels.",
        ],
        "B. Statistical reporting": [
            "Report macro and micro definitions.",
            "Separate internal all-label macro from external-evaluable macro.",
            "State 500 bootstrap iterations.",
            "Use patient-level bootstrap for NIH internal data.",
            "Use image-identifier-level bootstrap wording for external data.",
        ],
        "C. Dataset and label reporting": [
            "Report NIH 5000 images and 3795 patients.",
            "Report VinDr-derived external 1000 images.",
            "State NIH report-mined labels and VinDr radiologist-derived annotations.",
            "Mark Edema/Pneumonia external zero-positive.",
            "Add dataset-source/license details before submission.",
        ],
        "D. Calibration reporting": [
            "State NIH-only calibration fitting.",
            "State no external calibration fitting in primary analysis.",
            "Report Brier, ECE, MCE, slope, and intercept.",
            "Avoid global calibration-improvement claims.",
            "Interpret Platt/isotonic with slope/intercept caution.",
        ],
        "E. Grad-CAM++ reporting": [
            "Use failure-mode saliency language.",
            "Avoid localization-ground-truth claims.",
            "Avoid proof-of-reasoning language.",
            "Include TP/FP/FN/TN case categories.",
            "State qualitative-only role in captions and Discussion.",
        ],
        "F. Tables and figures": [
            "Use final12 main table structures.",
            "Keep NA values for externally unevaluable labels.",
            "Include AUPRC baseline/lift supplement.",
            "Mention selected public subset in captions.",
            "Check figure resolution and journal size limits.",
        ],
        "G. References": [
            "Verify all reference metadata in PubMed/Google Scholar.",
            "Verify DOI, volume, issue, and pages.",
            "Keep CLAIM/TRIPOD+AI/reporting references.",
            "Do not overuse radiologist-level papers as performance support.",
            "Add external-validation and dataset-shift CXR AI references if target journal expects them.",
        ],
        "H. Ethics and data availability": [
            "Confirm local IRB exemption/waiver wording.",
            "Add repository/archive link.",
            "State original dataset license requirements.",
            "Confirm checkpoint-sharing policy.",
            "Scrub private paths if files will be shared publicly.",
        ],
        "I. Journal formatting": [
            "Select target journal.",
            "Apply target reference style.",
            "Check word limit.",
            "Move oversized tables to supplement if needed.",
            "Prepare cover letter and reporting checklist.",
        ],
    }
    for category, values in items.items():
        for item in values:
            rows.append({"category": category, "check_item": item})
    return pd.DataFrame(rows)


def add_references(doc: Document) -> None:
    doc.add_heading("References", level=1)
    refs = [
        "1. Wang X, Peng Y, Lu L, Lu Z, Bagheri M, Summers RM. ChestX-ray8: Hospital-scale chest X-ray database and benchmarks on weakly-supervised classification and localization of common thorax diseases. CVPR. 2017.",
        "2. Nguyen HQ, Lam K, Le LT, et al. VinDr-CXR: An open dataset of chest X-rays with radiologist's annotations. Scientific Data. 2022;9:429.",
        "3. Irvin J, Rajpurkar P, Ko M, et al. CheXpert: A large chest radiograph dataset with uncertainty labels and expert comparison. AAAI. 2019.",
        "4. Johnson AEW, Pollard TJ, Berkowitz SJ, et al. MIMIC-CXR, a de-identified publicly available database of chest radiographs with free-text reports. Scientific Data. 2019;6:317.",
        "5. Bustos A, Pertusa A, Salinas JM, de la Iglesia-Vaya M. PadChest: A large chest x-ray image dataset with multi-label annotated reports. Medical Image Analysis. 2020;66:101797.",
        "6. Rajpurkar P, Irvin J, Zhu K, et al. CheXNet: Radiologist-level pneumonia detection on chest X-rays with deep learning. arXiv. 2017.",
        "7. Rajpurkar P, Irvin J, Ball RL, et al. Deep learning for chest radiograph diagnosis: A retrospective comparison of the CheXNeXt algorithm to practicing radiologists. PLoS Medicine. 2018;15:e1002686.",
        "8. Finlayson SG, Subbaswamy A, Singh K, et al. The clinician and dataset shift in artificial intelligence. New England Journal of Medicine. 2021;385:283-286.",
        "9. Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. ICML. 2017.",
        "10. Kuleshov V, Fenner N, Ermon S. Accurate uncertainties for deep learning using calibrated regression. ICML. 2018.",
        "11. Selvaraju RR, Cogswell M, Das A, Vedantam R, Parikh D, Batra D. Grad-CAM: Visual explanations from deep networks via gradient-based localization. ICCV. 2017.",
        "12. Chattopadhyay A, Sarkar A, Howlader P, Balasubramanian VN. Grad-CAM++: Generalized gradient-based visual explanations for deep convolutional networks. WACV. 2018.",
        "13. Huang G, Liu Z, van der Maaten L, Weinberger KQ. Densely connected convolutional networks. CVPR. 2017.",
        "14. He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition. CVPR. 2016.",
        "15. Tan M, Le QV. EfficientNet: Rethinking model scaling for convolutional neural networks. ICML. 2019.",
        "16. Mongan J, Moy L, Kahn CE Jr. Checklist for artificial intelligence in medical imaging (CLAIM): A guide for authors and reviewers. Radiology: Artificial Intelligence. 2020;2:e200029.",
        "17. Tejani AS, Klontzas ME, Gatti AA, et al. Checklist for artificial intelligence in medical imaging (CLAIM): 2024 update. Radiology: Artificial Intelligence. 2024.",
        "18. Collins GS, Dhiman P, Andaur Navarro CL, et al. Protocol for development of a reporting guideline for prediction model studies using machine learning: TRIPOD+AI. BMJ. 2024.",
        "19. Liu X, Rivera SC, Moher D, Calvert MJ, Denniston AK. Reporting guidelines for clinical trial reports for interventions involving artificial intelligence: CONSORT-AI extension. Nature Medicine. 2020.",
        "20. Rivera SC, Liu X, Chan AW, Denniston AK, Calvert MJ. Guidelines for clinical trial protocols for interventions involving artificial intelligence: SPIRIT-AI extension. BMJ. 2020.",
        "21. Vasey B, Nagendran M, Campbell B, et al. Reporting guideline for the early-stage clinical evaluation of decision support systems driven by artificial intelligence: DECIDE-AI. Nature Medicine. 2022.",
        "22. Bossuyt PM, Reitsma JB, Bruns DE, et al. STARD 2015: An updated list of essential items for reporting diagnostic accuracy studies. BMJ. 2015;351:h5527.",
        "23. Wolff RF, Moons KGM, Riley RD, et al. PROBAST: A tool to assess the risk of bias and applicability of prediction model studies. Annals of Internal Medicine. 2019;170:51-58.",
        "24. Nagendran M, Chen Y, Lovejoy CA, et al. Artificial intelligence versus clinicians: Systematic review of design, reporting standards, and claims. BMJ. 2020;368:m689.",
        "25. Davis J, Goadrich M. The relationship between precision-recall and ROC curves. ICML. 2006.",
        "26. Saito T, Rehmsmeier M. The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. PLoS ONE. 2015;10:e0118432.",
    ]
    for ref in refs:
        para(doc, ref)


def make_manuscript(path: Path, tables: dict[str, pd.DataFrame]) -> Path:
    sec = manuscript_sections()
    doc = Document()
    style_doc(doc)
    add_title(doc, sec["title"], f"Short title: {sec['short_title']}")
    para(doc, "Study type: retrospective public-dataset machine-learning validation study")

    doc.add_heading("Abstract", level=1)
    for p in sec["abstract"]:
        para(doc, p)
    para(doc, "Keywords: chest radiograph; public-subset validation; external-evaluable labels; calibration transfer; Grad-CAM++; NIH ChestX-ray14; VinDr-CXR")

    doc.add_heading("Introduction", level=1)
    for p in sec["introduction"]:
        para(doc, p)

    doc.add_heading("Methods", level=1)
    for heading, text in sec["methods"]:
        doc.add_heading(heading, level=2)
        para(doc, text)

    doc.add_heading("Results", level=1)
    for heading, text in sec["results"]:
        doc.add_heading(heading, level=2)
        para(doc, text)

    doc.add_heading("Discussion", level=1)
    for p in sec["discussion"]:
        para(doc, p)

    doc.add_heading("Limitations", level=1)
    for p in sec["limitations"]:
        para(doc, p)

    doc.add_heading("Conclusion", level=1)
    para(doc, sec["conclusion"])

    doc.add_heading("Data and Code Availability", level=1)
    para(doc, sec["data_code"])

    doc.add_heading("Ethics Statement", level=1)
    para(doc, sec["ethics"])

    doc.add_heading("Main Tables", level=1)
    add_df(doc, "Table 1. Dataset characteristics and split roles", tables["table1_dataset_characteristics_final12.csv"], ["dataset", "role", "split", "n_images", "n_patients_or_identifiers", "split_level", "used_for_training", "used_for_threshold_selection", "used_for_calibration_fitting", "used_for_final_evaluation"], max_rows=8)
    add_df(doc, "Table 2. Label harmonization and analysis role", tables["table2_label_harmonization_final12.csv"], ["harmonized_label", "NIH_terms", "VinDr_terms", "internal_analysis_role", "external_positive_cases", "external_evaluable", "external_claim_status", "caution"], max_rows=8)
    add_df(doc, "Table 3. DenseNet121 label-wise internal and external performance", tables["table3_densenet121_labelwise_performance_final12.csv"], ["label", "validation_threshold", "internal_prevalence", "external_prevalence", "internal_AUROC_95CI", "external_AUROC_95CI", "internal_AUPRC_95CI", "external_AUPRC_95CI", "external_status"], max_rows=8)
    add_df(doc, "Table 4. Macro-level architecture comparison", tables["table4_architecture_comparison_final12.csv"], ["model", "internal_all_label_macro_AUROC", "internal_all_label_macro_AUPRC", "internal_external_evaluable_macro_AUROC", "internal_external_evaluable_macro_AUPRC", "external_nonzero_prevalence_macro_AUROC", "external_nonzero_prevalence_macro_AUPRC"], max_rows=3)
    add_df(doc, "Table 5. Calibration metrics before and after NIH-fitted calibration", tables["table5_calibration_metrics_final12.csv"], ["dataset", "calibration_method", "Brier_score", "ECE", "MCE", "calibration_slope", "calibration_intercept", "external_fitting_used", "interpretation"], max_rows=8)
    add_df(doc, "Table 6. Descriptive subgroup and sensitivity analysis", tables["table6_subgroup_sensitivity_final12.csv"], ["dataset", "subgroup_variable", "subgroup", "n_images", "n_patients_or_identifiers", "macro_AUROC", "macro_AUPRC", "analysis_status"], max_rows=8)

    captions = figure_captions()
    doc.add_heading("Figures", level=1)
    fig_files = [
        ("Figure 1. Study workflow", "figure1_study_workflow_three_models.png"),
        ("Figure 2. DenseNet121 internal and external AUROC/AUPRC", "figure2_densenet121_internal_external_ci.png"),
        ("Figure 3. AUPRC baseline and lift", "figure3_prevalence_auprc_baseline_lift.png"),
        ("Figure 4. Calibration curves", "figure4_calibration_curves_upgraded.png"),
        ("Figure 5. Grad-CAM++ failure-mode panels", "figure5_gradcam_failure_modes_upgraded.png"),
        ("Figure 6. Internal-minus-external AUROC difference", "figure6_internal_minus_external_auroc.png"),
    ]
    for title, fname in fig_files:
        add_figure(doc, title, FIGS / fname, captions[title])

    doc.add_heading("Supplementary Material Overview", level=1)
    supp_items = [
        "Supplementary Table 1. ResNet50 label-wise performance.",
        "Supplementary Table 2. EfficientNet-B0 label-wise performance.",
        "Supplementary Table 3. Per-label thresholds from NIH validation.",
        "Supplementary Table 4. AUPRC prevalence baseline and lift.",
        "Supplementary Table 5. Per-label calibration metrics.",
        "Supplementary Table 6. Paired bootstrap model comparisons.",
        "Supplementary Table 7. Hyperparameters and environment.",
        "Supplementary Table 8. Reproducibility commands.",
        "Supplementary Table 9. Dataset-source and license summary.",
    ]
    bullets(doc, supp_items)
    add_references(doc)

    doc.save(path)
    shutil.copy2(path, OUTPUTS / path.name)
    return path


def make_report(path: Path, tables: dict[str, pd.DataFrame]) -> Path:
    sec = manuscript_sections()
    main_plan, supp_plan = table_plans()
    captions = figure_captions()
    overclaim = overclaim_table()
    checklist = checklist_table()
    doc = Document()
    style_doc(doc)
    add_title(doc, "12.docx Final Pre-submission Revision Report", "Systematic SCI manuscript strengthening and risk audit")

    doc.add_heading("Overall reviewer-style diagnosis", level=1)
    bullets(doc, [
        "The manuscript is now substantially closer to a defensible SCI submission because it clearly separates development, threshold selection, calibration fitting, internal testing, and external validation.",
        "The strongest contribution is the reproducible public-subset validation and calibration-transfer package rather than peak classification performance.",
        "The main residual risk is scope: this remains a selected public-subset analysis, not a full official-cohort or prospective clinical validation.",
        "The revised wording now avoids clinical-readiness, radiologist-level, full-cohort, and saliency-as-reasoning claims.",
    ])

    doc.add_heading("Priority revision summary", level=1)
    bullets(doc, [
        "Removed draft-style language such as word-count notes and journal-formatting placeholders.",
        "Changed external resampling language to image-identifier-level bootstrap approximation.",
        "Separated internal all-label macro results from external-evaluable macro results.",
        "Marked Edema and Pneumonia as zero-positive externally and not used for external discrimination claims.",
        "Rewrote calibration interpretation as method-dependent calibration transfer.",
        "Restricted Grad-CAM++ to qualitative failure-mode saliency.",
    ])

    doc.add_heading("Revised title options", level=1)
    add_df(doc, "Title options", pd.DataFrame(title_options()), ["title", "comment"])
    para(doc, f"Most recommended title: {sec['title']}")
    para(doc, f"Recommended short title: {sec['short_title']}")

    for heading, key in [
        ("Final revised Abstract", "abstract"),
        ("Final revised Introduction", "introduction"),
    ]:
        doc.add_heading(heading, level=1)
        for p in sec[key]:
            para(doc, p)

    doc.add_heading("Final revised Methods", level=1)
    for h, p in sec["methods"]:
        doc.add_heading(h, level=2)
        para(doc, p)

    doc.add_heading("Final revised Results", level=1)
    for h, p in sec["results"]:
        doc.add_heading(h, level=2)
        para(doc, p)

    doc.add_heading("Final revised Discussion", level=1)
    for p in sec["discussion"]:
        para(doc, p)

    doc.add_heading("Final revised Limitations", level=1)
    for p in sec["limitations"]:
        para(doc, p)

    doc.add_heading("Final revised Conclusion", level=1)
    para(doc, sec["conclusion"])

    doc.add_heading("Revised Data and Code Availability", level=1)
    para(doc, sec["data_code"])

    doc.add_heading("Revised Ethics Statement", level=1)
    para(doc, sec["ethics"])

    doc.add_heading("Main table restructuring plan", level=1)
    add_df(doc, "Main table plan", main_plan, ["table_name", "purpose", "recommended_columns", "placement", "interpretation_risk"], max_rows=6)

    doc.add_heading("Supplementary table restructuring plan", level=1)
    add_df(doc, "Supplementary table plan", supp_plan, ["table_name", "purpose", "recommended_columns", "placement", "interpretation_risk"], max_rows=9)

    doc.add_heading("Final figure captions", level=1)
    for title, cap in captions.items():
        doc.add_heading(title, level=2)
        para(doc, cap)

    doc.add_heading("Reference checking and expansion advice", level=1)
    bullets(doc, [
        "The current 26 references are adequate for a public-subset validation manuscript if the target journal accepts a concise reference list.",
        "Add 3 to 5 chest radiograph AI external-validation or dataset-shift papers if the target journal expects deeper external-validation context.",
        "Grad-CAM++ is present and should remain cited separately from Grad-CAM.",
        "Radiologist-level papers should be cited historically and cautiously, not used to imply comparable performance in this study.",
        "Optional added directions: external validation of CXR AI; medical AI calibration; public dataset label noise; weak-label CXR training; prevalence-aware PR analysis; AI reporting-guideline implementation; saliency limitations in medical imaging; dataset shift in clinical ML.",
        "Before submission, verify every reference in PubMed, Crossref, Google Scholar, or journal pages for author order, title, journal, year, volume, issue, pages, and DOI. Do not add DOI values unless verified; use [DOI TO BE VERIFIED] while drafting.",
    ])

    doc.add_heading("Terminology and overclaiming replacement table", level=1)
    add_df(doc, "Overclaiming risk phrase replacement table", overclaim, ["original_expression", "risk", "recommended_replacement"], max_rows=10)

    doc.add_heading("Final pre-submission checklist", level=1)
    add_df(doc, "Checklist", checklist, ["category", "check_item"], max_rows=len(checklist))

    doc.add_heading("SCI Q4 feasibility assessment", level=1)
    para(doc, "Feasibility is good for a Q4 medical imaging, radiology AI, or applied biomedical informatics journal if the manuscript is submitted as a transparent public-subset validation and calibration-transfer study. The reproducibility package, external-validation guardrails, and cautious claims are suitable strengths.")

    doc.add_heading("SCI Q3 feasibility assessment", level=1)
    para(doc, "Feasibility is possible but more sensitive to reviewer expectations. A Q3 target may request full official-cohort data, stronger external subgroup analysis, or deeper comparison with recent chest radiograph external-validation literature. Without prospective validation or full official cohorts, Q3 positioning should emphasize reproducibility, calibration transfer, and transparent reporting rather than performance leadership.")

    doc.add_heading("Remaining information that must be filled before submission", level=1)
    bullets(doc, [
        "Target journal and article type.",
        "Author names, affiliations, funding, acknowledgments, conflicts of interest, and corresponding author details.",
        "Public repository or archive link: [TO BE FILLED: repository or archive link].",
        "Local IRB exemption or waiver wording.",
        "Reference DOI/volume/issue/page verification.",
        "Final decision on whether model checkpoints may be shared publicly under dataset and institutional constraints.",
        "Any target-journal reporting checklist files, such as CLAIM/TRIPOD+AI/STARD-style appendices.",
    ])

    doc.save(path)
    shutil.copy2(path, OUTPUTS / path.name)
    return path


def qa_docx(paths: list[Path]) -> dict:
    banned_positive_patterns = [
        "Word-count note",
        "generated manuscript draft",
        "journal-specific formatting can be applied",
        "deployment-ready",
        "radiologist-level performance",
        "full official-cohort validation",
        "proof of model reasoning",
    ]
    summary = {}
    for path in paths:
        with zipfile.ZipFile(path) as zf:
            integrity = zf.testzip()
            media = [n for n in zf.namelist() if n.startswith("word/media/")]
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        summary[path.name] = {
            "path": str(path),
            "size_bytes": path.stat().st_size,
            "zip_integrity_error": integrity,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "inline_shape_count": len(doc.inline_shapes),
            "embedded_media_count": len(media),
            "banned_positive_phrases_found": [p for p in banned_positive_patterns if p in text],
            "contains_negative_clinical_readiness_disclaimer": "does not establish full official-cohort performance or clinical readiness" in text
            or "not establish full official-cohort performance, prospective validity, or clinical readiness" in text,
            "contains_external_identifier_bootstrap": "image-identifier-level" in text,
            "contains_repository_placeholder": "[TO BE FILLED: repository or archive link]" in text,
            "contains_gradcam_qualitative_only": "qualitative failure-mode" in text,
            "contains_zero_external_positive_statement": "zero external positives" in text or "zero positive cases" in text,
        }
    out = OUTPUTS / "docx_qa_summary_12_final.json"
    out.write_text(json.dumps(summary, indent=2), encoding="utf-8")
    return summary


def main() -> None:
    OUTPUTS.mkdir(exist_ok=True)
    tables = build_final12_tables()
    manuscript = make_manuscript(DESKTOP / "CXR_Manuscript_SCI_Submission_Final_12.docx", tables)
    report = make_report(DESKTOP / "CXR_12_PreSubmission_Revision_Report.docx", tables)
    summary = qa_docx([manuscript, report])
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()


