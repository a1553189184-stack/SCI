from __future__ import annotations

import json
import math
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
TABLES = ROOT / "tables_large"
FIGS = ROOT / "figures_large"
OUTPUTS = ROOT / "outputs_large"


def fmt(x: object, digits: int = 3) -> str:
    try:
        val = float(x)
    except Exception:
        return str(x)
    if not math.isfinite(val):
        return "NA"
    return f"{val:.{digits}f}"


def pct(x: object, digits: int = 1) -> str:
    try:
        val = float(x)
    except Exception:
        return str(x)
    if not math.isfinite(val):
        return "NA"
    return f"{100 * val:.{digits}f}%"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(15)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(10)


def para(doc: Document, text: str, style: str | None = None) -> None:
    doc.add_paragraph(text, style=style)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para(doc, item, style="List Bullet")


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.5 if len(table.columns) > 8 else 8)
                    if i == 0:
                        run.bold = True


def add_df(doc: Document, title: str, df: pd.DataFrame, cols: list[str], max_rows: int | None = None) -> None:
    doc.add_heading(title, level=2)
    use = df[cols].copy()
    if max_rows is not None:
        use = use.head(max_rows)
    table = doc.add_table(rows=1, cols=len(cols))
    for j, col in enumerate(cols):
        table.rows[0].cells[j].text = col
    for _, row in use.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(cols):
            val = row[col]
            if isinstance(val, float):
                cells[j].text = fmt(val)
            else:
                cells[j].text = str(val)
    style_table(table)


def image_for_docx(path: Path) -> Path:
    from PIL import Image

    out_dir = OUTPUTS / "docx_embedded_images_10_upgrade"
    out_dir.mkdir(parents=True, exist_ok=True)
    out = out_dir / f"{path.stem}.jpg"
    if out.exists() and out.stat().st_mtime >= path.stat().st_mtime:
        return out
    im = Image.open(path).convert("RGB")
    im.save(out, quality=94, optimize=True)
    return out


def add_figure(doc: Document, title: str, path: Path, caption: str, width: float = 6.4) -> None:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_for_docx(path)), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.style = "Caption"


def value_registry() -> dict:
    return json.loads((OUTPUTS / "manuscript_value_registry_10_upgrade.json").read_text(encoding="utf-8"))


def env_summary() -> dict:
    return json.loads((OUTPUTS / "training" / "hf_large_densenet121" / "environment_report.json").read_text(encoding="utf-8"))


def table(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


def manuscript_title() -> str:
    return "Cross-dataset Generalizability, Calibration Transfer, and Grad-CAM++ Failure-mode Analysis of Deep Learning Models for Multi-label Chest Radiograph Classification"


def add_references(doc: Document) -> None:
    doc.add_heading("References", level=1)
    refs = [
        "1. Wang X, Peng Y, Lu L, Lu Z, Bagheri M, Summers RM. ChestX-ray8: Hospital-scale chest X-ray database and benchmarks on weakly-supervised classification and localization of common thorax diseases. CVPR. 2017.",
        "2. Nguyen HQ, Lam K, Le LT, et al. VinDr-CXR: An open dataset of chest X-rays with radiologist's annotations. Scientific Data. 2022;9:429.",
        "3. Irvin J, Rajpurkar P, Ko M, et al. CheXpert: A large chest radiograph dataset with uncertainty labels and expert comparison. AAAI. 2019.",
        "4. Johnson AEW, Pollard TJ, Berkowitz SJ, et al. MIMIC-CXR, a de-identified publicly available database of chest radiographs with free-text reports. Scientific Data. 2019;6:317.",
        "5. Bustos A, Pertusa A, Salinas JM, de la Iglesia-Vaya M. PadChest: A large chest x-ray image dataset with multi-label annotated reports. Medical Image Analysis. 2020;66:101797.",
        "6. Rajpurkar P, Irvin J, Zhu K, et al. CheXNet: Radiologist-level pneumonia detection on chest X-rays with deep learning. arXiv. 2017.",
        "7. Rajpurkar P, Irvin J, Ball RL, et al. Deep learning for chest radiograph diagnosis: A retrospective comparison of the CheXNeXt algorithm to practicing radiologists. PLoS Medicine. 2018;15:e1002686.",
        "8. Finlayson SG, Subbaswamy A, Singh K, et al. The clinician and dataset shift in artificial intelligence. New England Journal of Medicine. 2021;385:283-286.",
        "9. Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. ICML. 2017.",
        "10. Kuleshov V, Fenner N, Ermon S. Accurate uncertainties for deep learning using calibrated regression. ICML. 2018.",
        "11. Selvaraju RR, Cogswell M, Das A, Vedantam R, Parikh D, Batra D. Grad-CAM: Visual explanations from deep networks via gradient-based localization. ICCV. 2017.",
        "12. Chattopadhyay A, Sarkar A, Howlader P, Balasubramanian VN. Grad-CAM++: Generalized gradient-based visual explanations for deep convolutional networks. WACV. 2018.",
        "13. Huang G, Liu Z, van der Maaten L, Weinberger KQ. Densely connected convolutional networks. CVPR. 2017.",
        "14. He K, Zhang X, Ren S, Sun J. Deep residual learning for image recognition. CVPR. 2016.",
        "15. Tan M, Le QV. EfficientNet: Rethinking model scaling for convolutional neural networks. ICML. 2019.",
        "16. Mongan J, Moy L, Kahn CE Jr. Checklist for artificial intelligence in medical imaging (CLAIM): A guide for authors and reviewers. Radiology: Artificial Intelligence. 2020;2:e200029.",
        "17. Tejani AS, Klontzas ME, Gatti AA, et al. Checklist for artificial intelligence in medical imaging (CLAIM): 2024 update. Radiology: Artificial Intelligence. 2024.",
        "18. Collins GS, Dhiman P, Andaur Navarro CL, et al. Protocol for development of a reporting guideline for prediction model studies using machine learning: TRIPOD+AI. BMJ. 2024.",
        "19. Liu X, Rivera SC, Moher D, Calvert MJ, Denniston AK. Reporting guidelines for clinical trial reports for interventions involving artificial intelligence: CONSORT-AI extension. Nature Medicine. 2020.",
        "20. Rivera SC, Liu X, Chan AW, Denniston AK, Calvert MJ. Guidelines for clinical trial protocols for interventions involving artificial intelligence: SPIRIT-AI extension. BMJ. 2020.",
        "21. Vasey B, Nagendran M, Campbell B, et al. Reporting guideline for the early-stage clinical evaluation of decision support systems driven by artificial intelligence: DECIDE-AI. Nature Medicine. 2022.",
        "22. Bossuyt PM, Reitsma JB, Bruns DE, et al. STARD 2015: An updated list of essential items for reporting diagnostic accuracy studies. BMJ. 2015;351:h5527.",
        "23. Wolff RF, Moons KGM, Riley RD, et al. PROBAST: A tool to assess the risk of bias and applicability of prediction model studies. Annals of Internal Medicine. 2019;170:51-58.",
        "24. Nagendran M, Chen Y, Lovejoy CA, et al. Artificial intelligence versus clinicians: Systematic review of design, reporting standards, and claims. BMJ. 2020;368:m689.",
        "25. Davis J, Goadrich M. The relationship between precision-recall and ROC curves. ICML. 2006.",
        "26. Saito T, Rehmsmeier M. The precision-recall plot is more informative than the ROC plot when evaluating binary classifiers on imbalanced datasets. PLoS ONE. 2015;10:e0118432.",
    ]
    for ref in refs:
        para(doc, ref)


def make_manuscript(path: Path) -> Path:
    reg = value_registry()
    env = env_summary()
    t1 = table("table1_dataset_characteristics_upgraded.csv")
    t2 = table("table2_label_harmonization_analysis_role_upgraded.csv")
    t3 = table("table3_densenet121_labelwise_performance.csv")
    t4 = table("table4_architecture_comparison_macro_upgraded.csv")
    t5 = table("table5_calibration_metrics_upgraded.csv")
    t6 = table("table6_subgroup_sensitivity_analysis.csv")
    auprc_lift = table("supplementary_table4_auprc_baseline_lift.csv")
    paired = table("supplementary_table6_paired_bootstrap_model_comparison.csv")

    d = reg["densenet121"]
    doc = Document()
    style_doc(doc)
    add_title(doc, manuscript_title(), "Retrospective public-dataset machine-learning validation study")
    para(doc, "Short title: NIH-to-VinDr public-subset chest radiograph validation")
    para(doc, "Word-count note: generated manuscript draft with embedded main tables and figures; journal-specific formatting can be applied after target selection.")

    doc.add_heading("Abstract", level=1)
    para(
        doc,
        "Background: Cross-dataset deployment of chest radiograph classifiers is limited by dataset shift, label-source heterogeneity, class imbalance, and imperfect probability calibration. Discrimination, calibration, and failure-mode explainability therefore need to be evaluated together rather than as isolated model-development outputs.",
    )
    para(
        doc,
        f"Methods: We built a reproducible PyTorch pipeline using a selected NIH ChestX-ray14 public parquet subset as the development dataset and a VinDr-CXR/VinBigData-derived public PNG subset as independent external validation. The NIH subset contained {reg['dataset']['nih_images']} images from {reg['dataset']['nih_patients']} patients and was split at patient level into training, validation, calibration, and internal-test partitions. The external subset contained {reg['dataset']['vindr_images']} images. DenseNet121 was prespecified as the main model; ResNet50 and EfficientNet-B0 were trained as architecture comparators. Per-label thresholds were selected only on NIH validation data, and temperature scaling, Platt scaling, and isotonic regression were fitted only on the NIH calibration split.",
    )
    para(
        doc,
        f"Results: DenseNet121 achieved internal all-label macro AUROC/AUPRC of {d['internal_all_label_macro_auroc']}/{d['internal_all_label_macro_auprc']}. For labels with nonzero external prevalence ({'; '.join(reg['external_evaluable_labels'])}), external macro AUROC/AUPRC were {d['external_nonzero_prevalence_macro_auroc']}/{d['external_nonzero_prevalence_macro_auprc']}; the corresponding NIH internal macro values on the same label set were {d['internal_external_evaluable_macro_auroc']}/{d['internal_external_evaluable_macro_auprc']}. ResNet50 and EfficientNet-B0 showed external nonzero-prevalence macro AUROC/AUPRC of {fmt(t4.loc[t4['model'].eq('ResNet50'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}/{fmt(t4.loc[t4['model'].eq('ResNet50'), 'external_nonzero_prevalence_labels_macro_auprc'].iloc[0])} and {fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}/{fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'external_nonzero_prevalence_labels_macro_auprc'].iloc[0])}, respectively. DenseNet121 internal uncalibrated macro Brier score, ECE, MCE, slope, and intercept were {reg['calibration']['internal_uncalibrated_brier']}, {reg['calibration']['internal_uncalibrated_ece']}, {reg['calibration']['internal_uncalibrated_mce']}, {reg['calibration']['internal_uncalibrated_slope']}, and {reg['calibration']['internal_uncalibrated_intercept']}. Grad-CAM++ panels were generated for internal and external true-positive, false-positive, false-negative, and true-negative cases.",
    )
    para(
        doc,
        "Conclusions: The pipeline provides an auditable public-dataset framework for assessing chest radiograph classifier transfer, calibration transfer, and saliency-based failure modes. The findings support cautious public-subset validation reporting, but they do not establish full official-cohort performance or clinical readiness.",
    )
    para(doc, "Keywords: chest radiograph; external validation; calibration; dataset shift; Grad-CAM++; DenseNet; VinDr-CXR; NIH ChestX-ray14")

    doc.add_heading("Introduction", level=1)
    intro_paras = [
        "Deep-learning systems for chest radiograph diagnosis have been widely evaluated on public datasets, yet their apparent performance can depend strongly on dataset origin, label generation, acquisition workflow, and disease prevalence. A model that ranks cases acceptably in one dataset may not preserve discrimination, decision thresholds, or probability calibration after transfer to another dataset.",
        "This problem is especially relevant for public chest radiograph resources. NIH ChestX-ray14 labels are derived from radiology reports, whereas VinDr-CXR provides radiologist annotations. These differences create a practical test of cross-dataset transportability but also introduce label-source heterogeneity that must be acknowledged when interpreting performance.",
        "AUROC alone is insufficient in this setting. AUPRC better reflects positive-label prevalence for imbalanced findings, thresholded metrics are sensitive to the dataset used for threshold selection, and probability calibration determines whether predicted risks correspond to observed event frequencies. Explainability maps can reveal plausible failure modes, but they should not be used as proof of radiological reasoning.",
        "We therefore designed a reproducible public-subset study to evaluate three linked questions: how a DenseNet121 chest radiograph classifier transfers from NIH ChestX-ray14 to a VinDr-derived external subset, whether NIH-fitted calibration procedures transfer to internal and external test data, and what Grad-CAM++ panels show in representative correct and incorrect decisions.",
        "DenseNet121 was chosen as the prespecified main model because it is a common chest radiograph baseline and has been used extensively in medical-image classification. ResNet50 and EfficientNet-B0 were added as architecture comparators to reduce single-model dependence while preserving the same data splits, thresholds, and external-validation rules.",
    ]
    for text in intro_paras:
        para(doc, text)

    doc.add_heading("Methods", level=1)
    methods = [
        ("Study design", "This was a retrospective public-dataset machine-learning validation study. The development dataset was a selected NIH ChestX-ray14 public parquet subset. Independent external validation used a VinDr-CXR/VinBigData-derived public PNG subset. All results in this manuscript were computed from saved prediction CSV files and metric tables."),
        ("Datasets", f"The NIH development subset included {reg['dataset']['nih_images']} images from {reg['dataset']['nih_patients']} patients. The external subset included {reg['dataset']['vindr_images']} images with image-level identifiers used as patient identifiers in the cached public subset. Dataset characteristics and split sizes are summarized in Table 1."),
        ("Labels", "The harmonized label set included Atelectasis, Cardiomegaly, Pleural Effusion, Pneumothorax, Consolidation, Edema, Pneumonia, and No Finding. External discrimination claims were restricted to labels with both positive and negative cases in the analyzed external subset. Edema and Pneumonia had zero external positives and were retained only for internal or sensitivity-context reporting."),
        ("Data splitting and leakage control", "NIH data were split at patient level into train, validation, calibration, and internal-test partitions. The validation split was used for threshold selection and checkpoint selection. The calibration split was used only for fitting calibration models. Internal-test and external data were not used for training, threshold optimization, checkpoint selection, or calibration fitting."),
        ("Preprocessing", "Images were converted to three-channel tensors, resized to 224 x 224 pixels, normalized using ImageNet mean and standard deviation, and augmented during training using horizontal flipping and small rotations. The same deterministic preprocessing was used for internal and external evaluation."),
        ("Models and training", f"DenseNet121 was trained as the main model, with ResNet50 and EfficientNet-B0 as comparators. All models used ImageNet pretrained torchvision weights, BCEWithLogitsLoss with positive class weights, AdamW optimization, mixed precision when CUDA was available, and early stopping/checkpoint selection using validation macro AUPRC. The environment used Python {env.get('python', 'unknown').split()[0]}, PyTorch {env.get('torch', 'unknown')}, CUDA availability {env.get('cuda_available', 'unknown')}, and GPU {env.get('gpu_name', 'unknown')}. Hyperparameters are listed in Supplementary Table 7."),
        ("Thresholds and evaluation", "Per-label thresholds were selected by maximizing F1 on the NIH validation split. These thresholds were then applied unchanged to the NIH internal-test split and VinDr external subset. Metrics included AUROC, AUPRC, sensitivity, specificity, F1-score, accuracy, PPV, NPV, and macro/micro averages. Patient-level bootstrap resampling with 500 iterations was used for 95% confidence intervals."),
        ("Calibration", "Temperature scaling, Platt scaling, and isotonic regression were fitted on the NIH calibration split only. Calibration transfer was evaluated on NIH internal-test and VinDr external predictions using Brier score, ECE, MCE, calibration slope, and calibration intercept. No calibration model was fitted on VinDr in the primary analysis."),
        ("Explainability", "Grad-CAM++ heatmaps were generated for selected true-positive, false-positive, false-negative, and true-negative examples from internal and external test data. These heatmaps were used for qualitative failure-mode review and not as evidence of clinical reasoning or localization validity."),
        ("Statistical analysis", "Macro discrimination was reported for all labels where calculable and separately for the external-evaluable label set. AUPRC baseline and lift were computed using observed label prevalence. Paired bootstrap model comparisons were reported descriptively rather than as confirmatory superiority tests."),
    ]
    for heading, text in methods:
        doc.add_heading(heading, level=2)
        para(doc, text)

    doc.add_heading("Results", level=1)
    result_paras = [
        f"The NIH development subset contained {reg['dataset']['nih_images']} images from {reg['dataset']['nih_patients']} patients. The patient-level split yielded {t1.loc[t1['split'].eq('train'), 'n_images'].iloc[0]} training images, {t1.loc[t1['split'].eq('validation'), 'n_images'].iloc[0]} validation images, {t1.loc[t1['split'].eq('calibration'), 'n_images'].iloc[0]} calibration images, and {t1.loc[t1['split'].eq('internal_test'), 'n_images'].iloc[0]} internal-test images. The external validation subset contained {reg['dataset']['vindr_images']} images.",
        f"DenseNet121 internal all-label macro AUROC/AUPRC were {d['internal_all_label_macro_auroc']}/{d['internal_all_label_macro_auprc']}. On the external-evaluable label set, NIH internal macro AUROC/AUPRC were {d['internal_external_evaluable_macro_auroc']}/{d['internal_external_evaluable_macro_auprc']}, whereas VinDr external macro AUROC/AUPRC were {d['external_nonzero_prevalence_macro_auroc']}/{d['external_nonzero_prevalence_macro_auprc']}. The higher external AUPRC should be read in the context of label prevalence and selected subset composition rather than as proof of broadly improved generalization.",
        f"Model comparison showed DenseNet121, ResNet50, and EfficientNet-B0 external nonzero-prevalence macro AUROC values of {fmt(t4.loc[t4['model'].eq('DenseNet121'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}, {fmt(t4.loc[t4['model'].eq('ResNet50'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}, and {fmt(t4.loc[t4['model'].eq('EfficientNet-B0'), 'external_nonzero_prevalence_labels_macro_auroc'].iloc[0])}, respectively. The paired bootstrap comparison favored DenseNet121 over ResNet50 for internal all-label AUROC, but most AUPRC and external-evaluable differences had confidence intervals crossing zero.",
        f"Calibration metrics indicated imperfect probability scaling before post-hoc calibration. DenseNet121 internal uncalibrated macro Brier score, ECE, MCE, calibration slope, and intercept were {reg['calibration']['internal_uncalibrated_brier']}, {reg['calibration']['internal_uncalibrated_ece']}, {reg['calibration']['internal_uncalibrated_mce']}, {reg['calibration']['internal_uncalibrated_slope']}, and {reg['calibration']['internal_uncalibrated_intercept']}. On external data, the lowest macro ECE among NIH-fitted methods was {reg['calibration']['best_external_ece']} using {reg['calibration']['best_external_ece_method']} scaling, but this should be interpreted with slope, intercept, and label prevalence.",
        "Grad-CAM++ examples were assembled for true-positive, false-positive, false-negative, and true-negative cases in NIH internal and VinDr external evaluation. Visual review suggested heterogeneous saliency patterns across errors, supporting their use as failure-mode illustrations rather than mechanistic explanations.",
    ]
    for text in result_paras:
        para(doc, text)

    doc.add_heading("Discussion", level=1)
    discussion = [
        "This study provides an end-to-end public-dataset validation package for chest radiograph classification under dataset shift. The design deliberately separates model selection, threshold tuning, calibration fitting, internal testing, and external validation, which reduces several common sources of optimistic bias.",
        "The external macro AUPRC exceeded internal AUPRC on the external-evaluable label set. This observation is compatible with differences in prevalence and case mix; it should not be framed as evidence that the model generalizes better externally than internally. The AUPRC baseline/lift analysis was therefore included to anchor precision-recall performance to observed prevalence.",
        "Calibration transfer remained an important limitation. Post-hoc methods fitted on the NIH calibration split could reduce some bin-wise error estimates, but calibration slope and intercept showed that probability-scale behavior is not fully solved by a single scalar or monotonic mapping. External calibration results are best interpreted as transfer diagnostics.",
        "The architecture-comparison results show that the overall conclusions are not unique to DenseNet121. However, DenseNet121 remains the prespecified main model, and comparator differences are reported descriptively because the analysis was not designed as a definitive architecture-superiority trial.",
        "Grad-CAM++ panels can support reviewer inspection of false positives and false negatives, but they are not localization ground truth and do not validate model reasoning. Their proper role in this manuscript is failure-mode illustration, not clinical explanation.",
        "The practical contribution of this work is the reproducibility package: all splits, thresholds, predictions, metric tables, figures, model checkpoints, logs, environment files, and random seeds are available locally in a structured package. This makes the work auditable and easier to revise if a journal requests additional sensitivity analyses.",
    ]
    for text in discussion:
        para(doc, text)

    doc.add_heading("Limitations", level=1)
    limitations = [
        "The analysis used executable public subsets, not the complete official NIH ChestX-ray14 or VinDr-CXR cohorts.",
        "NIH labels and VinDr labels come from different annotation processes, making label-source heterogeneity unavoidable.",
        "Edema and Pneumonia had zero external positives in the analyzed external subset and cannot support external discrimination claims.",
        "Images were resized to 224 x 224 pixels; the study should not be described as a full-resolution clinical-performance benchmark.",
        "The training schedule was intentionally compact and reproducible rather than tuned for maximum leaderboard performance.",
        "External metadata did not support reliable demographic subgroup analysis in the cached subset.",
        "Calibration was fitted internally and transferred externally; no external calibration fitting was performed in the primary analysis.",
        "Grad-CAM++ was qualitative and was not validated against expert localization in the main analysis.",
    ]
    bullets(doc, limitations)

    doc.add_heading("Conclusion", level=1)
    para(
        doc,
        "A reproducible NIH ChestX-ray14 to VinDr-derived public-subset pipeline was completed for multi-label chest radiograph classification. DenseNet121, ResNet50, and EfficientNet-B0 were evaluated with validation-derived thresholds, patient-level bootstrap confidence intervals, NIH-only calibration fitting, and Grad-CAM++ failure-mode review. The study supports transparent public-dataset reporting of discrimination, calibration transfer, and qualitative saliency behavior, while avoiding claims of clinical readiness.",
    )

    doc.add_heading("Data and Code Availability", level=1)
    para(
        doc,
        "The reproducibility package contains split CSVs, label mapping, prediction CSVs, calibration outputs, figure source data, model checkpoints, training logs, package versions, random seeds, GPU information, and rerun commands. Original images and labels derive from public NIH ChestX-ray14 and VinDr-CXR/VinBigData-derived resources; users should follow each dataset's source license and citation requirements.",
    )

    doc.add_heading("Ethics Statement", level=1)
    para(
        doc,
        "This study used de-identified public datasets and did not involve new patient recruitment, intervention, or direct patient contact. Local institutional review-board exemption requirements should be confirmed before journal submission.",
    )

    doc.add_heading("Main Tables", level=1)
    add_df(doc, "Table 1. Dataset characteristics and split roles", t1, ["dataset", "role", "split", "n_images", "n_patients", "patient_level_split"], max_rows=8)
    add_df(doc, "Table 2. Label harmonization and analysis role", t2, ["harmonized_label", "analysis_role", "external_positives_in_analyzed_subset", "external_evaluable_for_auc_pr", "caution"], max_rows=8)
    add_df(doc, "Table 3. DenseNet121 label-wise internal and external performance", t3, ["label", "validation_threshold", "internal_auroc_95ci", "external_auroc_95ci", "internal_auprc_95ci", "external_auprc_95ci", "external_analysis_status"], max_rows=8)
    add_df(doc, "Table 4. Macro architecture comparison", t4, ["model", "internal_all_labels_macro_auroc", "internal_all_labels_macro_auprc", "internal_external_evaluable_labels_macro_auroc", "internal_external_evaluable_labels_macro_auprc", "external_nonzero_prevalence_labels_macro_auroc", "external_nonzero_prevalence_labels_macro_auprc"], max_rows=3)
    add_df(doc, "Table 5. Calibration metrics before and after NIH-fitted calibration", t5, ["dataset", "calibration_method", "brier_score", "ece", "mce", "calibration_slope", "calibration_intercept", "external_fitting_used"], max_rows=8)
    add_df(doc, "Table 6. Subgroup and sensitivity analysis", t6, ["dataset", "subgroup_variable", "subgroup", "n_images", "all_label_macro_auroc", "all_label_macro_auprc", "analysis_status"], max_rows=8)

    doc.add_heading("Figures", level=1)
    add_figure(doc, "Figure 1. Study workflow", FIGS / "figure1_study_workflow_three_models.png", "The workflow separates NIH training, validation threshold selection, NIH-only calibration fitting, internal testing, and independent VinDr external validation.")
    add_figure(doc, "Figure 2. DenseNet121 internal and external AUROC/AUPRC", FIGS / "figure2_densenet121_internal_external_ci.png", "Label-wise AUROC and AUPRC with 95% bootstrap intervals. Edema and Pneumonia are marked NA externally because the analyzed external subset had zero positive cases.")
    add_figure(doc, "Figure 3. AUPRC baseline and lift", FIGS / "figure3_prevalence_auprc_baseline_lift.png", "Observed AUPRC is shown against prevalence baseline to contextualize precision-recall performance under label imbalance.")
    add_figure(doc, "Figure 4. Calibration curves", FIGS / "figure4_calibration_curves_upgraded.png", "Reliability diagrams for NIH internal and VinDr external predictions. Calibration models were fitted only on the NIH calibration split.")
    add_figure(doc, "Figure 5. Grad-CAM++ failure-mode panels", FIGS / "figure5_gradcam_failure_modes_upgraded.png", "Representative true-positive, false-positive, false-negative, and true-negative panels. Grad-CAM++ is used only for qualitative failure-mode analysis.")
    add_figure(doc, "Figure 6. Internal-minus-external AUROC difference", FIGS / "figure6_internal_minus_external_auroc.png", "Negative values indicate higher AUROC in the selected external subset; the plot should be interpreted with prevalence and label-source differences.")

    doc.add_heading("Supplementary Materials", level=1)
    supp_items = [
        "Supplementary Table 1: ResNet50 label-wise performance.",
        "Supplementary Table 2: EfficientNet-B0 label-wise performance.",
        "Supplementary Table 3: per-label thresholds from NIH validation.",
        "Supplementary Table 4: AUPRC prevalence baseline and lift.",
        "Supplementary Table 5: per-label calibration metrics.",
        "Supplementary Table 6: paired patient-level bootstrap model comparison.",
        "Supplementary Table 7: hyperparameters and environment.",
        "Supplementary Table 8: reproducibility commands.",
        "Supplementary Table 9: MedMNIST verification status, documenting that no MedMNIST metric is used as manuscript evidence.",
        "Supplementary Figure 1: ROC and PR curves for DenseNet121.",
        "Supplementary Figure 2: predicted probability distributions for DenseNet121.",
    ]
    bullets(doc, supp_items)
    add_df(doc, "Supplementary excerpt: paired bootstrap model comparison", paired, ["dataset", "label_set", "metric", "model_a", "model_b", "difference_a_minus_b", "ci_lower", "ci_upper", "ci_excludes_zero"], max_rows=12)
    add_df(doc, "Supplementary excerpt: AUPRC baseline/lift", auprc_lift, ["model", "dataset", "label", "positives", "auprc", "auprc_baseline_prevalence", "auprc_lift_over_prevalence", "status"], max_rows=16)

    add_references(doc)
    doc.save(path)
    OUTPUTS.mkdir(exist_ok=True)
    shutil.copy2(path, OUTPUTS / path.name)
    return path


def make_report(path: Path) -> Path:
    reg = value_registry()
    env = env_summary()
    t1 = table("table1_dataset_characteristics_upgraded.csv")
    t2 = table("table2_label_harmonization_analysis_role_upgraded.csv")
    t3 = table("table3_densenet121_labelwise_performance.csv")
    t4 = table("table4_architecture_comparison_macro_upgraded.csv")
    t5 = table("table5_calibration_metrics_upgraded.csv")
    paired = table("supplementary_table6_paired_bootstrap_model_comparison.csv")
    doc = Document()
    style_doc(doc)
    add_title(doc, "10.docx Systematic Upgrade Execution Report", "21 requested outputs completed from real experiment artifacts")

    doc.add_heading("1. Reviewer-style overall diagnosis", level=1)
    bullets(
        doc,
        [
            "Main strength: the project now has a complete public-dataset pipeline with patient-level NIH splitting, external validation, calibration transfer, comparator architectures, Grad-CAM++ failure-mode examples, and reproducibility packaging.",
            "Main risk: the analysis uses selected public subsets rather than full official cohorts; the manuscript must say public-subset validation and must not claim clinical readiness.",
            "Key interpretation risk: external AUPRC is higher than internal AUPRC, but this is driven by prevalence/case-mix/label-source differences and should not be framed as superior external generalization.",
            "Label risk: Edema and Pneumonia had zero external positives and must remain outside external primary discrimination claims.",
        ],
    )

    doc.add_heading("2. Revised paper positioning", level=1)
    para(doc, "Recommended positioning: a reproducible public-subset external-validation and calibration-transfer study for multi-label chest radiograph classification, using NIH ChestX-ray14 as development data and VinDr-derived data as independent external validation.")
    para(doc, "Avoided positioning: full official-cohort benchmark, prospective clinical validation, or clinical decision-support readiness study.")

    doc.add_heading("3. Recommended titles", level=1)
    bullets(
        doc,
        [
            manuscript_title(),
            "Reproducible Public-subset External Validation of Deep Learning Chest Radiograph Classifiers from NIH ChestX-ray14 to VinDr-CXR",
            "Discrimination, Calibration Transfer, and Saliency Failure Modes in Public Chest Radiograph Classification",
            "From NIH ChestX-ray14 to VinDr-CXR: External Validation and Calibration of Multi-label Chest Radiograph Classifiers",
            "Dataset Shift, Probability Calibration, and Grad-CAM++ Failure Modes in Public Chest X-ray AI Models",
        ],
    )

    doc.add_heading("4. Rewritten Abstract", level=1)
    para(doc, "The abstract in the upgraded manuscript has been fully rewritten with structured Background, Methods, Results, and Conclusions. It uses actual sample sizes, macro metrics, calibration values, external-evaluable labels, and no placeholders.")

    doc.add_heading("5. Rewritten Introduction", level=1)
    para(doc, "The Introduction now frames the study around dataset shift, label-source heterogeneity, AUPRC under imbalance, calibration transfer, and cautious saliency use.")

    doc.add_heading("6. Rewritten Methods", level=1)
    bullets(
        doc,
        [
            f"Environment: Python {env.get('python', 'unknown').split()[0]}, PyTorch {env.get('torch', 'unknown')}, CUDA {env.get('torch_cuda', 'unknown')}, GPU {env.get('gpu_name', 'unknown')}.",
            "NIH split: patient-level train/validation/calibration/internal-test.",
            "Thresholds: selected only from NIH validation.",
            "Calibration: fitted only on NIH calibration split.",
            "External validation: VinDr external predictions were used only after model, threshold, and calibration choices were fixed.",
        ],
    )

    doc.add_heading("7. Rewritten Results", level=1)
    d = reg["densenet121"]
    para(
        doc,
        f"DenseNet121 internal all-label macro AUROC/AUPRC were {d['internal_all_label_macro_auroc']}/{d['internal_all_label_macro_auprc']}; external nonzero-prevalence macro AUROC/AUPRC were {d['external_nonzero_prevalence_macro_auroc']}/{d['external_nonzero_prevalence_macro_auprc']}. Internal values on the same external-evaluable label set were {d['internal_external_evaluable_macro_auroc']}/{d['internal_external_evaluable_macro_auprc']}.",
    )
    para(
        doc,
        f"Calibration baseline: internal uncalibrated macro Brier/ECE/MCE/slope/intercept = {reg['calibration']['internal_uncalibrated_brier']}/{reg['calibration']['internal_uncalibrated_ece']}/{reg['calibration']['internal_uncalibrated_mce']}/{reg['calibration']['internal_uncalibrated_slope']}/{reg['calibration']['internal_uncalibrated_intercept']}.",
    )

    doc.add_heading("8. Rewritten Discussion", level=1)
    para(doc, "The Discussion now emphasizes that external performance must be interpreted with prevalence and label-source differences, calibration transfer remains imperfect, and Grad-CAM++ is an exploratory failure-mode tool.")

    doc.add_heading("9. Rewritten Limitations", level=1)
    bullets(
        doc,
        [
            "Selected public subsets rather than full official cohorts.",
            "Weak-label NIH source versus radiologist-annotated VinDr source.",
            "Zero-positive Edema and Pneumonia externally.",
            "224 x 224 preprocessing, not a full-resolution official-cohort benchmark.",
            "No prospective or clinical workflow validation.",
        ],
    )

    doc.add_heading("10. Rewritten Conclusion", level=1)
    para(doc, "The Conclusion now states that the project provides a transparent reproducible validation package, not evidence for clinical deployment.")

    doc.add_heading("11. Data/Code Availability", level=1)
    para(doc, "The reproducibility package includes split CSVs, label_map.yaml, prediction CSVs, calibration outputs, figure source data, checkpoints, logs, package versions, random seeds, and GPU information.")

    doc.add_heading("12. Ethics Statement", level=1)
    para(doc, "The manuscript states that the study used de-identified public datasets and requires local confirmation of IRB exemption or waiver documentation before submission.")

    doc.add_heading("13. Main table design", level=1)
    add_df(doc, "Main Table 1", t1, ["dataset", "role", "split", "n_images", "n_patients"], max_rows=8)
    add_df(doc, "Main Table 2", t2, ["harmonized_label", "analysis_role", "external_positives_in_analyzed_subset", "external_evaluable_for_auc_pr"], max_rows=8)
    add_df(doc, "Main Table 3", t3, ["label", "internal_auroc_95ci", "external_auroc_95ci", "internal_auprc_95ci", "external_auprc_95ci", "external_analysis_status"], max_rows=8)
    add_df(doc, "Main Table 4", t4, ["model", "internal_all_labels_macro_auroc", "internal_all_labels_macro_auprc", "external_nonzero_prevalence_labels_macro_auroc", "external_nonzero_prevalence_labels_macro_auprc"], max_rows=3)
    add_df(doc, "Main Table 5", t5, ["dataset", "calibration_method", "brier_score", "ece", "mce", "calibration_slope", "calibration_intercept"], max_rows=8)

    doc.add_heading("14. Supplementary table design", level=1)
    bullets(
        doc,
        [
            "Supplementary Table 1: ResNet50 label-wise performance.",
            "Supplementary Table 2: EfficientNet-B0 label-wise performance.",
            "Supplementary Table 3: NIH-validation thresholds.",
            "Supplementary Table 4: AUPRC baseline and lift.",
            "Supplementary Table 5: per-label calibration metrics.",
            "Supplementary Table 6: paired bootstrap model comparison.",
            "Supplementary Table 7: hyperparameters.",
            "Supplementary Table 8: environment and rerun commands.",
            "Supplementary Table 9: MedMNIST status, explicitly not used as evidence.",
        ],
    )
    add_df(doc, "Supplementary Table 6 excerpt", paired, ["dataset", "label_set", "metric", "model_a", "model_b", "difference_a_minus_b", "ci_lower", "ci_upper", "ci_excludes_zero"], max_rows=12)

    doc.add_heading("15. Figure design and captions", level=1)
    bullets(
        doc,
        [
            "Figure 1: workflow with explicit train/validation/calibration/internal/external separation.",
            "Figure 2: DenseNet121 label-wise AUROC/AUPRC with NA markings for externally unevaluable labels.",
            "Figure 3: prevalence baseline and AUPRC lift.",
            "Figure 4: calibration curves, specifying NIH-only calibration fitting.",
            "Figure 5: Grad-CAM++ failure-mode panels with caution against clinical reasoning claims.",
            "Figure 6: internal-minus-external AUROC difference.",
        ],
    )
    add_figure(doc, "Figure preview 1", FIGS / "figure1_study_workflow_three_models.png", "Workflow preview.", width=6.2)
    add_figure(doc, "Figure preview 2", FIGS / "figure2_densenet121_internal_external_ci.png", "Discrimination preview.", width=6.2)
    add_figure(doc, "Figure preview 3", FIGS / "figure3_prevalence_auprc_baseline_lift.png", "AUPRC lift preview.", width=6.2)

    doc.add_heading("16. Mandatory statistics list", level=1)
    bullets(
        doc,
        [
            "External-evaluable label macro AUROC/AUPRC: completed.",
            "Internal all-label macro and internal external-evaluable macro: completed.",
            "External macro restricted to nonzero-prevalence labels: completed.",
            "AUPRC baseline and lift: completed.",
            "Calibration Brier/ECE/MCE/slope/intercept: completed.",
            "Per-label thresholds: completed.",
            "Paired patient-level bootstrap model comparison: completed.",
            "ROC/PR curves and probability distributions: completed as supplementary figures.",
        ],
    )

    doc.add_heading("17. Reference expansion advice", level=1)
    bullets(
        doc,
        [
            "Keep dataset citations: NIH ChestX-ray14/ChestX-ray8, VinDr-CXR, CheXpert, MIMIC-CXR, PadChest.",
            "Keep model citations: DenseNet, ResNet, EfficientNet.",
            "Keep evaluation citations: calibration, precision-recall under imbalance, dataset shift.",
            "Keep reporting citations: CLAIM 2024, TRIPOD+AI, CONSORT-AI/SPIRIT-AI, DECIDE-AI, STARD, PROBAST.",
        ],
    )

    doc.add_heading("18. Pre-submission checklist", level=1)
    bullets(
        doc,
        [
            "No synthetic metrics in tables.",
            "No external test tuning.",
            "No VinDr calibration fitting in primary analysis.",
            "All external Edema/Pneumonia claims removed or marked unevaluable.",
            "Figures have source CSVs where numeric.",
            "Prediction CSVs are retained.",
            "Random seed and environment logged.",
            "Clinical claims are cautious.",
        ],
    )

    doc.add_heading("19. SCI Q4 feasibility", level=1)
    para(doc, "Feasibility: good for a methods/validation-oriented Q4 medical imaging or applied AI journal if the manuscript remains honest about public-subset scope and emphasizes reproducibility, calibration, and external-validation guardrails.")

    doc.add_heading("20. SCI Q3 feasibility", level=1)
    para(doc, "Feasibility: possible but more demanding. The strongest path is to position the work as a reproducible validation and calibration-transfer package, add full official-cohort data if feasible, and strengthen reporting-guideline compliance. Without full official cohorts or prospective validation, Q3 acceptance depends on venue tolerance for public-subset methodological studies.")

    doc.add_heading("21. Remaining data/file list", level=1)
    bullets(
        doc,
        [
            "Optional stronger evidence: full official NIH ChestX-ray14 image release and full official VinDr-CXR release.",
            "Optional stronger analysis: external demographic subgroup labels if reliable metadata are available.",
            "Optional stronger explainability: expert localization review or verified bounding-box overlap on externally positive findings.",
            "Administrative before submission: target journal formatting, author affiliations, funding, conflict-of-interest statement, and local IRB/exemption wording.",
        ],
    )

    doc.add_heading("Final reviewer-risk audit", level=1)
    bullets(doc, reg["risk_guardrails"])
    para(doc, "Audit result: no data leakage was identified in the documented NIH patient-level split; VinDr was not used for tuning, retraining, model selection, or calibration fitting; MedMNIST is not used as evidence; Grad-CAM++ language is restricted to qualitative failure-mode analysis.")

    doc.save(path)
    OUTPUTS.mkdir(exist_ok=True)
    shutil.copy2(path, OUTPUTS / path.name)
    return path


def main() -> None:
    manuscript = make_manuscript(DESKTOP / "CXR_Manuscript_SCI_Submission_Upgraded_10.docx")
    report = make_report(DESKTOP / "CXR_10_Systematic_Upgrade_Report.docx")
    print(manuscript)
    print(report)


if __name__ == "__main__":
    main()


