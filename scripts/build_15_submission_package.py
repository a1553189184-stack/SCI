from __future__ import annotations

import json
import math
import shutil
import sys
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from build_12_final_submission_docs import (  # noqa: E402
    FIGS,
    OUTPUTS,
    TABLES,
    add_df,
    add_figure,
    add_references,
    add_title,
    bullets,
    checklist_table,
    env_report,
    figure_captions,
    fmt,
    manuscript_sections,
    overclaim_table,
    para,
    registry,
    style_doc,
    table,
    table_plans,
    title_options,
)

DESKTOP = Path.home() / "Desktop"


def word_count(text: str) -> int:
    import re

    return len(re.findall(r"[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?", text))


def build_final15_main_tables() -> dict[str, pd.DataFrame]:
    t1 = table("table1_dataset_characteristics_final12.csv")
    t2 = table("table2_label_harmonization_final12.csv")
    t3 = table("table3_densenet121_labelwise_performance_final12.csv")
    t4 = table("table4_architecture_comparison_final12.csv")
    t5 = table("table5_calibration_metrics_final12.csv")
    t6 = table("table6_subgroup_sensitivity_final12.csv")

    def use_phrase(row: pd.Series) -> str:
        uses = []
        if row["used_for_training"] == "yes":
            uses.append("training")
        if row["used_for_threshold_selection"] == "yes":
            uses.append("threshold/checkpoint selection")
        if row["used_for_calibration_fitting"] == "yes":
            uses.append("calibration fitting")
        if row["used_for_final_evaluation"] == "yes":
            uses.append("final evaluation")
        return "; ".join(uses) if uses else "not directly analyzed as a split"

    table1 = pd.DataFrame(
        {
            "dataset": t1["dataset"],
            "role": t1["role"],
            "split": t1["split"],
            "n_images": t1["n_images"],
            "n_patients_or_identifiers": t1["n_patients_or_identifiers"],
            "split_or_identifier_level": t1["split_level"],
            "analysis_use": t1.apply(use_phrase, axis=1),
        }
    )

    table2 = pd.DataFrame(
        {
            "label": t2["harmonized_label"],
            "NIH_term": t2["NIH_terms"],
            "VinDr_term": t2["VinDr_terms"],
            "external_positive_cases": t2["external_positive_cases"],
            "external_status": t2["external_claim_status"],
            "caution": t2["caution"].fillna(""),
        }
    )

    table3 = pd.DataFrame(
        {
            "label": t3["label"],
            "threshold": t3["validation_threshold"],
            "prevalence_internal_external": t3.apply(
                lambda r: f"{fmt(r['internal_prevalence'])}/{fmt(r['external_prevalence'])}", axis=1
            ),
            "internal_AUROC_95CI": t3["internal_AUROC_95CI"],
            "external_AUROC_95CI": t3["external_AUROC_95CI"],
            "internal_AUPRC_95CI": t3["internal_AUPRC_95CI"],
            "external_AUPRC_95CI": t3["external_AUPRC_95CI"],
            "external_status": t3["external_status"],
        }
    )

    table4 = t4.copy()
    table4.columns = [
        "model",
        "internal_all_AUROC",
        "internal_all_AUPRC",
        "internal_ext_eval_AUROC",
        "internal_ext_eval_AUPRC",
        "external_nonzero_AUROC",
        "external_nonzero_AUPRC",
    ]

    table5 = pd.DataFrame(
        {
            "dataset": t5["dataset"],
            "method": t5["calibration_method"],
            "Brier": t5["Brier_score"],
            "ECE": t5["ECE"],
            "MCE": t5["MCE"],
            "slope": t5["calibration_slope"],
            "intercept": t5["calibration_intercept"],
            "interpretation": t5["interpretation"],
        }
    )

    table6_supp = t6.copy()
    out = {
        "table1_dataset_characteristics_main_final15.csv": table1,
        "table2_external_evaluable_labels_main_final15.csv": table2,
        "table3_densenet121_main_performance_final15.csv": table3,
        "table4_architecture_macro_main_final15.csv": table4,
        "table5_calibration_transfer_main_final15.csv": table5,
        "supplementary_table_descriptive_subgroup_final15.csv": table6_supp,
    }
    for name, df in out.items():
        df.to_csv(TABLES / name, index=False)
    return out


def declarations_text() -> dict[str, str]:
    return {
        "Funding": "No specific funding was received for this study.",
        "Conflict of interest": "The authors declare no competing interests.",
        "Data availability": "The study used public NIH ChestX-ray14 and VinDr-CXR/VinBigData-derived resources. Split CSVs, label mapping, prediction CSVs, calibration outputs, figure source data, logs, package versions, random seeds, GPU/environment information, and rerun commands are available at [TO BE FILLED: repository or archive link]. Original image and label files remain subject to the source dataset licenses and citation requirements.",
        "Code availability": "The analysis code, configuration files, and reproducibility scripts are available at [TO BE FILLED: repository or archive link]. Model checkpoint sharing should be confirmed against dataset and institutional requirements before public release.",
        "Ethics approval": "This retrospective study used de-identified public datasets and involved no new patient recruitment, intervention, or direct patient contact. Institutional review-board exemption or waiver requirements should be confirmed locally before submission. Approval number: [TO BE FILLED if required].",
        "Consent to participate": "Not applicable because this study used de-identified public datasets and did not involve new participant recruitment.",
        "Consent for publication": "Not applicable because no identifiable individual-level information is reported.",
        "Author contributions": "Conceptualization: [TO BE FILLED]. Data curation: [TO BE FILLED]. Methodology: [TO BE FILLED]. Software: [TO BE FILLED]. Formal analysis: [TO BE FILLED]. Writing, original draft: [TO BE FILLED]. Writing, review and editing: [TO BE FILLED]. Supervision: [TO BE FILLED]. All authors approved the final manuscript: [TO BE CONFIRMED].",
        "Acknowledgements": "[TO BE FILLED] or The authors thank the maintainers of the NIH ChestX-ray14 and VinDr-CXR public datasets.",
        "Use of AI tools": "AI-assisted tools were used for language editing, formatting support, and manuscript-organization assistance. All scientific claims, numerical results, code outputs, tables, and figures were checked by the authors against the underlying experiment artifacts. No AI tool was used to generate or modify patient data.",
        "Clinical trial registration": "Not applicable. This was a retrospective public-dataset machine-learning validation study and did not involve an interventional clinical trial.",
        "Patient and public involvement": "Not applicable. Patients or members of the public were not involved in the design, conduct, reporting, or dissemination planning of this public-dataset analysis.",
    }


def cover_letter_text() -> list[str]:
    title = title_options()[0]["title"]
    return [
        "[TO BE FILLED: Date]",
        "",
        "Dear [TO BE FILLED: Editor Name],",
        "",
        f"We submit our manuscript entitled \"{title}\" for consideration in [TO BE FILLED: Journal Name].",
        "",
        "This manuscript reports a retrospective public-dataset machine-learning validation study for multi-label chest radiograph classification. We evaluated a DenseNet121 main model with ResNet50 and EfficientNet-B0 architecture comparators using a selected NIH ChestX-ray14 public subset for development and a selected VinDr-CXR/VinBigData-derived public subset for external validation.",
        "",
        "The study emphasizes several safeguards that are important for medical imaging AI reporting: NIH patient-level splitting, validation-only threshold selection, NIH-only calibration fitting, external-evaluable label restrictions, method-dependent calibration-transfer interpretation, and Grad-CAM++ used only for qualitative failure-mode review. The manuscript does not claim clinical readiness or full official-cohort validation.",
        "",
        "We believe the work may be of interest to readers of [TO BE FILLED: Journal Name] because it provides an auditable public-subset framework for evaluating discrimination, prevalence-aware precision-recall behavior, calibration transfer, and saliency-based failure modes under dataset shift.",
        "",
        "The manuscript has not been published previously and is not under consideration elsewhere [TO BE CONFIRMED]. All authors have approved this submission [TO BE CONFIRMED]. The authors declare no competing interests. Data and code availability details are provided in the manuscript, with the repository or archive link to be added before submission: [TO BE FILLED: repository or archive link].",
        "",
        "Thank you for considering our manuscript.",
        "",
        "Sincerely,",
        "[TO BE FILLED: Corresponding Author Name]",
        "[TO BE FILLED: Affiliation]",
        "[TO BE FILLED: Email]",
    ]


def highlights() -> list[str]:
    return [
        "Public-subset validation from NIH ChestX-ray14 to VinDr-derived data",
        "Thresholds and calibration were fitted only on NIH development splits",
        "External claims were restricted to nonzero-prevalence labels",
        "Calibration transfer showed method-dependent effects",
        "Grad-CAM++ supported qualitative failure-mode review",
    ]


def graphical_abstract_plan() -> list[str]:
    return [
        "Left panel: NIH ChestX-ray14 public subset, 5,000 images from 3,795 patients, patient-level train/validation/calibration/internal-test split.",
        "Middle panel: DenseNet121 main model with ResNet50 and EfficientNet-B0 comparators; validation-derived thresholds; NIH-only temperature, Platt, and isotonic calibration fitting.",
        "Right panel: selected VinDr-derived external subset, 1,000 images; external-evaluable labels only; discrimination, calibration transfer, and Grad-CAM++ failure-mode review.",
        "Bottom caution band: public-subset validation, not full official-cohort validation and not clinical readiness.",
    ]


def save_graphical_abstract(path: Path) -> None:
    import matplotlib.pyplot as plt
    from matplotlib.patches import FancyBboxPatch

    fig, ax = plt.subplots(figsize=(11.5, 5.0), dpi=180)
    ax.axis("off")
    ax.text(0.50, 0.91, "Public-subset chest radiograph validation workflow", ha="center", va="center", fontsize=13, fontweight="bold", color="#0F172A")
    boxes = [
        ("NIH ChestX-ray14 public subset\n5,000 images, 3,795 patients\nPatient-level train/validation/\ncalibration/internal-test split", 0.04, 0.40, "#E8F1F2"),
        ("Model development\nDenseNet121 main model\nResNet50 and EfficientNet-B0 comparators\nValidation-derived thresholds\nNIH-only calibration fitting", 0.36, 0.40, "#F4F1DE"),
        ("Selected VinDr-derived external subset\n1,000 images\nExternal-evaluable labels only\nDiscrimination, calibration transfer,\nGrad-CAM++ failure modes", 0.68, 0.40, "#FDECEF"),
    ]
    for text, x, y, color in boxes:
        patch = FancyBboxPatch((x, y), 0.27, 0.42, boxstyle="round,pad=0.018,rounding_size=0.015", fc=color, ec="#334155", lw=1.0)
        ax.add_patch(patch)
        ax.text(x + 0.135, y + 0.21, text, ha="center", va="center", fontsize=9)
    ax.annotate("", xy=(0.36, 0.61), xytext=(0.31, 0.61), arrowprops=dict(arrowstyle="->", lw=1.2, color="#334155"))
    ax.annotate("", xy=(0.68, 0.61), xytext=(0.63, 0.61), arrowprops=dict(arrowstyle="->", lw=1.2, color="#334155"))
    caution = FancyBboxPatch((0.14, 0.12), 0.72, 0.12, boxstyle="round,pad=0.015,rounding_size=0.012", fc="#FFF7ED", ec="#B45309", lw=1.0)
    ax.add_patch(caution)
    ax.text(0.50, 0.18, "Caution: public-subset validation, not full official-cohort validation and not clinical readiness", ha="center", va="center", fontsize=9, color="#7C2D12")
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=300)
    fig.savefig(path.with_suffix(".pdf"))
    plt.close(fig)


def export_figure_tiffs() -> pd.DataFrame:
    from PIL import Image

    out_dir = FIGS / "final15_tiff_300dpi"
    out_dir.mkdir(parents=True, exist_ok=True)
    rows = []
    names = [
        "figure1_study_workflow_three_models.png",
        "figure2_densenet121_internal_external_ci.png",
        "figure3_prevalence_auprc_baseline_lift.png",
        "figure4_calibration_curves_upgraded.png",
        "figure5_gradcam_failure_modes_upgraded.png",
        "figure6_internal_minus_external_auroc.png",
        "graphical_abstract_plan_final15.png",
    ]
    for name in names:
        src = FIGS / name
        if not src.exists():
            continue
        im = Image.open(src).convert("RGB")
        dst = out_dir / f"{src.stem}.tiff"
        im.save(dst, dpi=(300, 300), compression="tiff_lzw")
        rows.append(
            {
                "figure": src.stem,
                "source_png": str(src),
                "tiff_300dpi": str(dst),
                "width_px": im.width,
                "height_px": im.height,
                "dpi": "300",
                "format_recommendation": "Submit TIFF if required by journal; PNG/PDF may be acceptable for review upload.",
            }
        )
    df = pd.DataFrame(rows)
    df.to_csv(TABLES / "figure_quality_and_exports_final15.csv", index=False)
    return df


def risk_table() -> pd.DataFrame:
    rows = [
        ("Selected public subset", "Reviewers may question generalizability to full cohorts.", "Explicitly stated throughout.", "Keep public-subset wording in title, abstract, methods, captions.", "High", "selected public external subset"),
        ("Incomplete external labels", "Some NIH labels are not externally evaluable.", "External-evaluable set defined.", "Keep macro summaries restricted to nonzero-prevalence labels.", "High", "external-evaluable labels"),
        ("Edema/Pneumonia zero positives", "External AUROC/AUPRC are not calculable.", "Marked as not used for external claims.", "Keep NA values and explain in Table 2/Figure 2.", "High", "not used for external discrimination claims"),
        ("External patient identifiers limited", "Affects bootstrap interpretation.", "Image identifiers used.", "State image-identifier-level bootstrap approximation.", "High", "image-identifier-level bootstrap approximation"),
        ("External bootstrap approximation", "CIs may be optimistic if images are correlated.", "Disclosed.", "Retain as limitation.", "Medium", "interpreted as an approximation"),
        ("AUPRC prevalence sensitivity", "External AUPRC can rise with prevalence/case mix.", "AUPRC lift table generated.", "Discuss prevalence baseline.", "High", "prevalence-aware interpretation"),
        ("Calibration transfer instability", "ECE alone can be misleading.", "Brier/ECE/MCE/slope/intercept reported.", "Avoid global improvement claims.", "High", "method-dependent calibration effects"),
        ("Grad-CAM++ overinterpretation", "Saliency is not proof of reasoning.", "Qualitative-only wording used.", "Keep limitation in caption and discussion.", "High", "qualitative failure-mode review"),
        ("No prospective validation", "Limits clinical inference.", "Acknowledged.", "Do not claim deployment readiness.", "High", "does not establish clinical readiness"),
        ("No reader study", "Cannot compare with radiologists.", "Acknowledged.", "Avoid radiologist-level claims.", "Medium", "not compared with readers"),
        ("Table formatting", "Wide tables may be hard to fit.", "Final15 compact tables created.", "Move subgroup table to supplement if journal space is limited.", "Medium", "compact main table with explanatory footnote"),
        ("Figure quality", "Journal may require 300 dpi TIFF.", "300 dpi TIFF exports generated.", "Upload TIFF or PDF according to journal.", "Medium", "300 dpi figure export"),
        ("Repository link missing", "Data/code availability incomplete.", "Placeholder retained.", "Fill repository or archive link.", "High", "[TO BE FILLED: repository or archive link]"),
        ("Reference format not finalized", "Submission can be delayed or rejected at technical check.", "Reference advice generated.", "Verify DOI and apply Vancouver style.", "Medium", "[DOI TO BE VERIFIED] where needed"),
    ]
    return pd.DataFrame(rows, columns=["Risk item", "Why it matters", "Current status", "Required action", "Severity", "Suggested wording"])


def reference_advice() -> list[str]:
    return [
        "The current 26-reference structure is reasonable for an SCI Q4 submission if formatted correctly and checked against journal requirements.",
        "For an SCI Q3 target, add 5 to 10 references focused on chest radiograph external validation, clinical AI generalization, dataset shift, and calibration in medical imaging.",
        "Reduce reliance on arXiv where peer-reviewed versions exist. CheXNet may remain as historical context, but radiologist-level language should not frame this study's claims.",
        "Keep Grad-CAM++ as a separate citation from Grad-CAM.",
        "Verify all DOI values, volume/issue/page numbers, and author order before submission. Do not invent DOI values; use [DOI TO BE VERIFIED] until checked.",
        "Apply the target journal's Vancouver style or reference-manager export format before upload.",
        "Convert conference-paper formatting to journal-required style only after target journal selection.",
    ]


def submission_checklist() -> pd.DataFrame:
    rows = []
    categories = {
        "A. Scientific claims": ["No clinical-readiness claim.", "No full official-cohort claim.", "No radiologist-level claim.", "External AUPRC interpreted with prevalence.", "Comparator claims remain descriptive."],
        "B. Dataset reporting": ["NIH subset size stated.", "VinDr-derived subset size stated.", "Selected public subset stated.", "External identifiers described cautiously.", "Dataset licenses acknowledged."],
        "C. Statistical reporting": ["Threshold selection from NIH validation only.", "NIH bootstrap patient-level.", "External bootstrap image-identifier-level approximation.", "Macro rules stated.", "Paired bootstrap described as descriptive."],
        "D. Calibration reporting": ["NIH-only calibration fitting stated.", "No external calibration fitting stated.", "Brier/ECE/MCE/slope/intercept reported.", "Temperature scaling not overstated.", "Platt/isotonic interpreted cautiously."],
        "E. Explainability reporting": ["Grad-CAM++ qualitative only.", "No localization-ground-truth claim.", "No reasoning-proof claim.", "TP/FP/FN/TN categories labeled.", "Additional panels placed in supplement if needed."],
        "F. Tables": ["Use compact final15 main tables.", "Move subgroup table to supplement if space constrained.", "Use short column labels.", "Add footnotes for split usage and NA labels.", "Avoid replacing NA with zero."],
        "G. Figures": ["Export 300 dpi TIFF if required.", "Check figure fonts and legends.", "Figure 2 marks Edema/Pneumonia NA.", "Figure 4 states NIH-only calibration fitting.", "Figure 6 includes zero reference line and negative-value explanation."],
        "H. Supplementary materials": ["Include supplementary methods.", "Include comparator label-wise tables.", "Include thresholds.", "Include AUPRC lift.", "Include ROC/PR and probability distribution figures."],
        "I. References": ["Verify DOI values.", "Apply Vancouver style.", "Add external-validation references for Q3 target.", "Reduce arXiv reliance where possible.", "Check reporting-guideline citations."],
        "J. Declarations": ["Fill author contributions.", "Fill affiliations and corresponding author.", "Confirm conflicts.", "Confirm funding.", "Confirm IRB wording."],
        "K. Cover letter": ["Fill journal name.", "Fill editor name.", "Confirm no prior publication.", "Confirm no concurrent submission.", "Confirm all authors approved."],
        "L. Journal formatting": ["Check abstract word limit.", "Check manuscript word limit.", "Check table/figure upload format.", "Prepare reporting checklist.", "Scrub private paths before public release."],
    }
    for category, items in categories.items():
        for item in items:
            rows.append({"category": category, "check_item": item})
    return pd.DataFrame(rows)


def title_page_rows(abstract_words: int, main_word_count: int) -> list[tuple[str, str]]:
    title = title_options()[0]["title"]
    return [
        ("Full title", title),
        ("Short title", "Public-subset chest radiograph validation"),
        ("Article type", "Original research; retrospective public-dataset machine-learning validation study"),
        ("Author names", "[TO BE FILLED]"),
        ("Affiliations", "[TO BE FILLED]"),
        ("Corresponding author", "[TO BE FILLED]"),
        ("Email", "[TO BE FILLED]"),
        ("ORCID", "[TO BE FILLED]"),
        ("Word count", f"[TO BE FILLED according to target-journal rules; current manuscript text approximately {main_word_count} words]"),
        ("Abstract word count", f"{abstract_words} words, verify against target-journal counting rules"),
        ("Number of tables", "5 main tables; descriptive subgroup table recommended for supplement"),
        ("Number of figures", "6 main figures"),
        ("Number of references", "26 current references, final formatting to be verified"),
        ("Keywords", "chest radiograph; public-subset validation; calibration transfer; external-evaluable labels; Grad-CAM++; NIH ChestX-ray14; VinDr-CXR"),
        ("Running title", "Public-subset chest radiograph validation"),
    ]


def add_title_page(doc: Document, abstract_words: int, main_word_count: int) -> None:
    add_title(doc, title_options()[0]["title"], "Title page")
    rows = title_page_rows(abstract_words, main_word_count)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Item"
    tbl.rows[0].cells[1].text = "Content"
    for key, value in rows:
        cells = tbl.add_row().cells
        cells[0].text = key
        cells[1].text = value
    from build_12_final_submission_docs import style_table

    style_table(tbl)
    doc.add_page_break()


def make_final_manuscript(path: Path, main_tables: dict[str, pd.DataFrame]) -> Path:
    sec = manuscript_sections()
    abstract_words = word_count(" ".join(sec["abstract"]))
    main_words = word_count(" ".join(sec["abstract"] + sec["introduction"] + [p for _, p in sec["methods"]] + [p for _, p in sec["results"]] + sec["discussion"] + sec["limitations"] + [sec["conclusion"], sec["data_code"], sec["ethics"]]))
    doc = Document()
    style_doc(doc)
    add_title_page(doc, abstract_words, main_words)
    add_title(doc, title_options()[0]["title"], "Manuscript")
    para(doc, "Study type: retrospective public-dataset machine-learning validation study")
    doc.add_heading("Abstract", level=1)
    for p in sec["abstract"]:
        para(doc, p)
    para(doc, "Keywords: chest radiograph; public-subset validation; external-evaluable labels; calibration transfer; Grad-CAM++; NIH ChestX-ray14; VinDr-CXR")
    doc.add_heading("Introduction", level=1)
    for p in sec["introduction"]:
        para(doc, p)
    doc.add_heading("Methods", level=1)
    for h, p in sec["methods"]:
        doc.add_heading(h, level=2)
        para(doc, p)
    doc.add_heading("Results", level=1)
    for h, p in sec["results"]:
        doc.add_heading(h, level=2)
        para(doc, p)
    doc.add_heading("Discussion", level=1)
    for p in sec["discussion"]:
        para(doc, p)
    doc.add_heading("Limitations", level=1)
    for p in sec["limitations"]:
        para(doc, p)
    doc.add_heading("Conclusion", level=1)
    para(doc, sec["conclusion"])
    doc.add_heading("Data and Code Availability", level=1)
    para(doc, sec["data_code"])
    doc.add_heading("Declarations", level=1)
    for key, value in declarations_text().items():
        doc.add_heading(key, level=2)
        para(doc, value)
    doc.add_heading("Main Tables", level=1)
    add_df(doc, "Table 1. Dataset characteristics", main_tables["table1_dataset_characteristics_main_final15.csv"], list(main_tables["table1_dataset_characteristics_main_final15.csv"].columns), max_rows=8)
    para(doc, "Footnote: NIH splits were patient-level. The selected VinDr-derived external subset used image identifiers because reliable patient-level identifiers were not available in the cached subset.")
    add_df(doc, "Table 2. Externally evaluable labels and harmonization", main_tables["table2_external_evaluable_labels_main_final15.csv"], list(main_tables["table2_external_evaluable_labels_main_final15.csv"].columns), max_rows=8)
    para(doc, "Footnote: Edema and Pneumonia had zero external positives and were not used for external discrimination claims.")
    add_df(doc, "Table 3. DenseNet121 main performance", main_tables["table3_densenet121_main_performance_final15.csv"], list(main_tables["table3_densenet121_main_performance_final15.csv"].columns), max_rows=8)
    para(doc, "Footnote: Thresholds were selected on the NIH validation split only. NA indicates not calculable because the external subset had no positive cases.")
    add_df(doc, "Table 4. Architecture-level macro comparison", main_tables["table4_architecture_macro_main_final15.csv"], list(main_tables["table4_architecture_macro_main_final15.csv"].columns), max_rows=3)
    para(doc, "Footnote: External macro values are restricted to nonzero-prevalence external labels.")
    add_df(doc, "Table 5. Calibration transfer metrics", main_tables["table5_calibration_transfer_main_final15.csv"], list(main_tables["table5_calibration_transfer_main_final15.csv"].columns), max_rows=8)
    para(doc, "Footnote: Calibration methods were fitted only on the NIH calibration split; no external calibration fitting was performed.")
    captions = figure_captions()
    doc.add_heading("Figures", level=1)
    for title, fname in [
        ("Figure 1. Study workflow", "figure1_study_workflow_three_models.png"),
        ("Figure 2. DenseNet121 internal and external AUROC/AUPRC", "figure2_densenet121_internal_external_ci.png"),
        ("Figure 3. AUPRC baseline and lift", "figure3_prevalence_auprc_baseline_lift.png"),
        ("Figure 4. Calibration curves", "figure4_calibration_curves_upgraded.png"),
        ("Figure 5. Grad-CAM++ failure-mode panels", "figure5_gradcam_failure_modes_upgraded.png"),
        ("Figure 6. Internal-minus-external AUROC difference", "figure6_internal_minus_external_auroc.png"),
    ]:
        add_figure(doc, title, FIGS / fname, captions[title])
    add_references(doc)
    doc.save(path)
    shutil.copy2(path, OUTPUTS / path.name)
    return path


def make_cover_letter_doc(path: Path) -> Path:
    doc = Document()
    style_doc(doc)
    for line in cover_letter_text():
        if line:
            para(doc, line)
        else:
            doc.add_paragraph("")
    doc.save(path)
    shutil.copy2(path, OUTPUTS / path.name)
    return path


def make_supplementary_doc(path: Path) -> Path:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Supplementary Materials Organization", "Public-subset chest radiograph validation")
    para(doc, "Supplementary materials should begin with a brief note that all supplementary analyses use the same fixed NIH development splits, validation-derived thresholds, NIH-only calibration fitting, and selected VinDr-derived external subset described in the main manuscript.")
    doc.add_heading("Supplementary Methods", level=1)
    bullets(doc, [
        "Dataset extraction and public-subset construction.",
        "Label harmonization and external-evaluable label definition.",
        "Patient-level NIH split verification and external identifier limitations.",
        "Threshold selection, bootstrap resampling, and calibration fitting details.",
        "Grad-CAM++ case-selection and failure-mode review procedure.",
    ])
    doc.add_heading("Supplementary Tables", level=1)
    for title, desc in [
        ("Supplementary Table 1. ResNet50 label-wise performance", "Comparator label-wise internal and external metrics."),
        ("Supplementary Table 2. EfficientNet-B0 label-wise performance", "Comparator label-wise internal and external metrics."),
        ("Supplementary Table 3. Per-label thresholds from NIH validation", "Threshold provenance and validation F1 values."),
        ("Supplementary Table 4. AUPRC prevalence baseline and lift", "Prevalence-aware precision-recall interpretation."),
        ("Supplementary Table 5. Per-label calibration metrics", "Label-level Brier, ECE, MCE, slope, and intercept."),
        ("Supplementary Table 6. Paired bootstrap model comparisons", "Descriptive architecture comparisons."),
        ("Supplementary Table 7. Hyperparameters and environment", "Training and software environment details."),
        ("Supplementary Table 8. Reproducibility commands", "Commands used to rerun the pipeline."),
        ("Supplementary Table 9. Dataset-source and license summary", "Dataset source, access, and citation notes."),
        ("Supplementary Table 10. Descriptive subgroup analysis", "Moved from main text unless target journal requests all tables in the manuscript."),
    ]:
        doc.add_heading(title, level=2)
        para(doc, desc)
    doc.add_heading("Supplementary Figures", level=1)
    for title, cap in [
        ("Supplementary Figure 1. DenseNet121 ROC and precision-recall curves", "Per-label ROC and precision-recall curves for NIH internal and selected VinDr-derived external evaluation. External labels with zero positive cases are omitted from curve panels."),
        ("Supplementary Figure 2. DenseNet121 predicted probability distributions", "Predicted probability distributions by dataset and label, used to inspect calibration and domain-shift behavior."),
        ("Supplementary Figure 3. Additional Grad-CAM++ failure-mode panels", "Additional qualitative true-positive, false-positive, false-negative, and true-negative examples. These panels are not localization ground truth."),
    ]:
        doc.add_heading(title, level=2)
        para(doc, cap)
    doc.save(path)
    shutil.copy2(path, OUTPUTS / path.name)
    return path


def make_submission_materials_doc(path: Path, figure_quality: pd.DataFrame) -> Path:
    sec = manuscript_sections()
    main_plan, supp_plan = table_plans()
    risk = risk_table()
    checklist = submission_checklist()
    main_tables = build_final15_main_tables()
    doc = Document()
    style_doc(doc)
    add_title(doc, "Final Submission Package", "Title page, declarations, cover letter, highlights, graphical abstract plan, supplementary organization, risk table, and checklist")
    doc.add_heading("Overall final submission diagnosis", level=1)
    bullets(doc, [
        "The manuscript is ready for target-journal formatting after administrative placeholders are filled.",
        "Scientific claims are appropriately bounded to selected public-subset validation.",
        "External validation, calibration transfer, and Grad-CAM++ language are cautious.",
        "NIH uncertainty intervals are described as patient-level bootstrap estimates, whereas external uncertainty intervals are described as image-identifier-level bootstrap approximations.",
        "Main remaining blockers are target journal selection, author/declaration details, repository link, reference formatting, and IRB wording.",
    ])
    doc.add_heading("Final title options", level=1)
    add_df(doc, "Title options", pd.DataFrame(title_options()), ["title", "comment"], max_rows=5)
    doc.add_heading("Final abstract", level=1)
    for p in sec["abstract"]:
        para(doc, p)
    doc.add_heading("Final polished manuscript sections", level=1)
    para(doc, "The polished manuscript sections are included in CXR_Manuscript_SCI_Submission_Final_15.docx.")
    doc.add_heading("Main table revision plan", level=1)
    add_df(doc, "Main table plan", main_plan, ["table_name", "purpose", "recommended_columns", "placement", "interpretation_risk"], max_rows=6)
    doc.add_heading("Final figure captions", level=1)
    for title, cap in figure_captions().items():
        doc.add_heading(title, level=2)
        para(doc, cap)
    doc.add_heading("Supplementary material structure", level=1)
    add_df(doc, "Supplementary table plan", supp_plan, ["table_name", "purpose", "recommended_columns", "placement", "interpretation_risk"], max_rows=9)
    doc.add_heading("Title page template", level=1)
    abstract_words = word_count(" ".join(sec["abstract"]))
    main_words = word_count(" ".join(sec["abstract"] + sec["introduction"] + [p for _, p in sec["methods"]] + [p for _, p in sec["results"]] + sec["discussion"] + sec["limitations"] + [sec["conclusion"]]))
    add_df(doc, "Title page items", pd.DataFrame(title_page_rows(abstract_words, main_words), columns=["Item", "Content"]), ["Item", "Content"], max_rows=15)
    doc.add_heading("Declarations", level=1)
    for k, v in declarations_text().items():
        doc.add_heading(k, level=2)
        para(doc, v)
    doc.add_heading("Cover letter", level=1)
    for line in cover_letter_text():
        para(doc, line) if line else doc.add_paragraph("")
    doc.add_heading("Highlights", level=1)
    bullets(doc, highlights())
    doc.add_heading("Graphical abstract plan", level=1)
    bullets(doc, graphical_abstract_plan())
    add_figure(doc, "Graphical abstract schematic", FIGS / "graphical_abstract_plan_final15.png", "Draft graphical abstract layout. Text and styling can be adapted to the target journal's graphical abstract specifications.")
    doc.add_heading("Pre-submission risk table", level=1)
    add_df(doc, "Risk table", risk, ["Risk item", "Why it matters", "Current status", "Required action", "Severity", "Suggested wording"], max_rows=len(risk))
    doc.add_heading("Reference final-check advice", level=1)
    bullets(doc, reference_advice())
    doc.add_heading("Final submission checklist", level=1)
    add_df(doc, "Checklist", checklist, ["category", "check_item"], max_rows=len(checklist))
    doc.add_heading("Figure export status", level=1)
    add_df(doc, "300 dpi exports", figure_quality, ["figure", "width_px", "height_px", "dpi", "format_recommendation"], max_rows=len(figure_quality))
    doc.add_heading("SCI Q4 feasibility assessment", level=1)
    para(doc, "SCI Q4 feasibility is good after placeholders are filled and references are formatted. The strongest pitch is a reproducible public-subset validation and calibration-transfer package with clear reporting safeguards.")
    doc.add_heading("SCI Q3 feasibility assessment", level=1)
    para(doc, "SCI Q3 feasibility is possible but less secure. A Q3 journal may request full official-cohort validation, more external-validation literature, stronger external subgroup analysis, or a clearer repository release. Positioning should emphasize transparency, calibration transfer, and reproducibility rather than performance leadership.")
    doc.add_heading("Remaining items to fill before submission", level=1)
    bullets(doc, [
        "Target journal and editor name.",
        "Author names, affiliations, corresponding author, email, and ORCID.",
        "Repository or archive link.",
        "IRB exemption or waiver wording.",
        "Reference DOI/volume/issue/page verification.",
        "Final conflict-of-interest and funding confirmation.",
        "Journal-specific abstract and manuscript word-count rules.",
    ])
    doc.save(path)
    shutil.copy2(path, OUTPUTS / path.name)
    return path


def qa_docx(paths: list[Path]) -> dict:
    banned = [
        "generated manuscript draft",
        "journal-specific formatting can be applied",
        "deployment-ready",
        "radiologist-level performance",
        "proof of model reasoning",
        "Calibration improved reliability",
        "The model generalized better externally",
    ]
    summary = {}
    for path in paths:
        with zipfile.ZipFile(path) as zf:
            integrity = zf.testzip()
            media = [n for n in zf.namelist() if n.startswith("word/media/")]
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        summary[path.name] = {
            "path": str(path),
            "size_bytes": path.stat().st_size,
            "zip_integrity_error": integrity,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "inline_shape_count": len(doc.inline_shapes),
            "embedded_media_count": len(media),
            "banned_positive_phrases_found": [p for p in banned if p in text],
            "contains_negative_full_cohort_disclaimer": "not full official-cohort validation" in text
            or "does not establish full official-cohort performance" in text
            or "not full official-cohort validation" in text,
            "contains_selected_public_subset": "selected public" in text or "public-subset" in text,
            "contains_external_identifier_bootstrap": "image-identifier-level" in text,
            "contains_repository_placeholder": "[TO BE FILLED: repository or archive link]" in text,
            "contains_gradcam_qualitative": "qualitative failure-mode" in text,
            "contains_zero_external_positive": "zero external positives" in text or "zero positive cases" in text,
        }
    out = OUTPUTS / "docx_qa_summary_15_final.json"
    out.write_text(json.dumps(summary, indent=2), encoding="utf-8")
    return summary


def write_manifest(paths: list[Path], main_tables: dict[str, pd.DataFrame], figure_quality: pd.DataFrame) -> None:
    manifest = {
        "created_from": "15.docx final submission package request",
        "docx_outputs": [str(p) for p in paths],
        "main_tables": {name: str(TABLES / name) for name in main_tables},
        "figure_quality_csv": str(TABLES / "figure_quality_and_exports_final15.csv"),
        "figure_tiffs": figure_quality["tiff_300dpi"].tolist() if not figure_quality.empty else [],
        "graphical_abstract": str(FIGS / "graphical_abstract_plan_final15.png"),
        "risk_table_csv": str(TABLES / "pre_submission_risk_table_final15.csv"),
        "checklist_csv": str(TABLES / "final_submission_checklist_final15.csv"),
    }
    (OUTPUTS / "submission_package_manifest_final15.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def main() -> None:
    OUTPUTS.mkdir(exist_ok=True)
    main_tables = build_final15_main_tables()
    save_graphical_abstract(FIGS / "graphical_abstract_plan_final15.png")
    figure_quality = export_figure_tiffs()
    risk_table().to_csv(TABLES / "pre_submission_risk_table_final15.csv", index=False)
    submission_checklist().to_csv(TABLES / "final_submission_checklist_final15.csv", index=False)
    overclaim_table().to_csv(TABLES / "overclaiming_replacement_table_final15.csv", index=False)
    manuscript = make_final_manuscript(DESKTOP / "CXR_Manuscript_SCI_Submission_Final_15.docx", main_tables)
    materials = make_submission_materials_doc(DESKTOP / "CXR_Final_Submission_Package_15.docx", figure_quality)
    cover = make_cover_letter_doc(DESKTOP / "CXR_Cover_Letter_Final_15.docx")
    supp = make_supplementary_doc(DESKTOP / "CXR_Supplementary_Materials_Organization_Final_15.docx")
    paths = [manuscript, materials, cover, supp]
    summary = qa_docx(paths)
    write_manifest(paths, main_tables, figure_quality)
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()


